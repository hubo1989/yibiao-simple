"""Word 格式规范导出服务 - 从 Markdown 内容生成规范的 Word 文档"""
from __future__ import annotations

import io
import logging
import os
import re
import tempfile
import uuid
from dataclasses import dataclass, field
from typing import Any, Dict, List, Optional, Tuple

import docx
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

logger = logging.getLogger(__name__)

# ============ GB/T 9704 标准排版常量 ============
FONT_SIZE_XIAO_ER = Pt(18)   # 小二
FONT_SIZE_SAN = Pt(16)       # 三号
FONT_SIZE_XIAO_SAN = Pt(15)  # 小三
FONT_SIZE_SI = Pt(14)        # 四号
FONT_SIZE_XIAO_SI = Pt(12)   # 小四号
FONT_SIZE_WU = Pt(10.5)      # 五号

FONT_SONGTI = "宋体"
FONT_HEITI = "黑体"
FONT_FANGSONG = "仿宋"

# 表头颜色
TABLE_HEADER_BG = RGBColor(0x1E, 0x3A, 0x5F)   # 深蓝 #1E3A5F
TABLE_ALT_ROW_BG = RGBColor(0xF0, 0xF4, 0xFA)   # 交替行浅蓝


@dataclass
class WordFormatConfig:
    """Word 格式配置"""
    # 字体
    body_font: str = FONT_SONGTI
    body_size_pt: float = 12.0        # 小四号
    h1_font: str = FONT_SONGTI
    h1_size_pt: float = 18.0          # 小二
    h2_font: str = FONT_HEITI
    h2_size_pt: float = 16.0          # 三号
    h3_font: str = FONT_HEITI
    h3_size_pt: float = 15.0          # 小三
    table_font: str = FONT_SONGTI
    table_size_pt: float = 10.5       # 五号
    # 间距
    line_spacing_pt: float = 22.0     # 1.5 倍行距（五号字 * 1.5 ≈ 22pt）
    first_indent_chars: int = 2
    h1_before: float = 12.0
    h1_after: float = 6.0
    h2_before: float = 8.0
    h2_after: float = 4.0
    h3_before: float = 6.0
    h3_after: float = 3.0
    # 页边距（mm）
    margin_top_mm: float = 25.0
    margin_bottom_mm: float = 25.0
    margin_left_mm: float = 30.0
    margin_right_mm: float = 25.0
    # 页眉
    header_text: str = "{project_name}"
    # 页脚水印
    footer_text: str = "严格保密"
    # 目录
    show_toc: bool = True
    toc_title: str = "目  录"
    toc_levels: int = 3


@dataclass
class TableFormatConfig:
    """表格格式配置"""
    header_bg_color: RGBColor = field(default_factory=lambda: TABLE_HEADER_BG)
    header_font_color: RGBColor = field(default_factory=lambda: RGBColor(0xFF, 0xFF, 0xFF))
    alt_row_bg_color: RGBColor = field(default_factory=lambda: TABLE_ALT_ROW_BG)
    font_name: str = FONT_SONGTI
    font_size: Pt = FONT_SIZE_WU


def _parse_word_config_from_template(format_config: Dict[str, Any] | None) -> WordFormatConfig:
    """从 ExportTemplate.format_config JSON 解析 WordFormatConfig"""
    if not format_config:
        return WordFormatConfig()

    cfg = WordFormatConfig()
    font = format_config.get("font", {})
    spacing = format_config.get("spacing", {})
    margin = format_config.get("margin", {})
    page = format_config.get("page", {})
    toc = format_config.get("toc", {})

    if font.get("body_font"):
        cfg.body_font = font["body_font"]
    if font.get("body_size"):
        cfg.body_size_pt = float(font["body_size"])
    if font.get("h1_font"):
        cfg.h1_font = font["h1_font"]
    if font.get("h1_size"):
        cfg.h1_size_pt = float(font["h1_size"])
    if font.get("h2_font"):
        cfg.h2_font = font["h2_font"]
    if font.get("h2_size"):
        cfg.h2_size_pt = float(font["h2_size"])
    if font.get("h3_font"):
        cfg.h3_font = font["h3_font"]
    if font.get("h3_size"):
        cfg.h3_size_pt = float(font["h3_size"])
    if font.get("table_font"):
        cfg.table_font = font["table_font"]
    if font.get("table_size"):
        cfg.table_size_pt = float(font["table_size"])

    if spacing.get("line_spacing_pt"):
        cfg.line_spacing_pt = float(spacing["line_spacing_pt"])
    if spacing.get("first_indent_chars") is not None:
        cfg.first_indent_chars = int(spacing["first_indent_chars"])

    if margin.get("top") is not None:
        cfg.margin_top_mm = float(margin["top"])
    if margin.get("bottom") is not None:
        cfg.margin_bottom_mm = float(margin["bottom"])
    if margin.get("left") is not None:
        cfg.margin_left_mm = float(margin["left"])
    if margin.get("right") is not None:
        cfg.margin_right_mm = float(margin["right"])

    if page.get("header_text"):
        cfg.header_text = page["header_text"]
    if toc.get("show_toc") is not None:
        cfg.show_toc = bool(toc["show_toc"])
    if toc.get("toc_title"):
        cfg.toc_title = toc["toc_title"]

    return cfg


def _set_run_font(
    run: docx.text.run.Run,
    font_name: str = FONT_SONGTI,
    font_size: Optional[Pt] = None,
    bold: bool = False,
    color: Optional[RGBColor] = None,
) -> None:
    """设置 run 的字体"""
    run.font.name = font_name
    if font_size:
        run.font.size = font_size
    run.bold = bold
    if color:
        run.font.color.rgb = color
    # 设置 eastAsia 字体
    if run._element.rPr is None:
        run._element._add_rPr()
    if run._element.rPr.rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        run._element.rPr.insert(0, rFonts)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def _apply_paragraph_format(
    paragraph: docx.text.paragraph.Paragraph,
    cfg: WordFormatConfig,
    line_spacing_pt: Optional[float] = None,
    first_indent: bool = True,
) -> None:
    """应用段落格式"""
    pf = paragraph.paragraph_format
    ls = line_spacing_pt or cfg.line_spacing_pt
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(ls)
    if first_indent and cfg.first_indent_chars > 0:
        pf.first_line_indent = Pt(cfg.body_size_pt * cfg.first_indent_chars)
    else:
        pf.first_line_indent = Pt(0)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)


def _setup_section(doc: Document, cfg: WordFormatConfig) -> None:
    """设置页面页边距"""
    for section in doc.sections:
        section.top_margin = Cm(cfg.margin_top_mm / 10)
        section.bottom_margin = Cm(cfg.margin_bottom_mm / 10)
        section.left_margin = Cm(cfg.margin_left_mm / 10)
        section.right_margin = Cm(cfg.margin_right_mm / 10)


def _setup_header_footer(doc: Document, project_name: str, cfg: WordFormatConfig) -> None:
    """设置页眉和页脚"""
    for section in doc.sections:
        # 页眉：项目名称居中
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hr = hp.add_run(cfg.header_text.replace("{project_name}", project_name or ""))
        _set_run_font(hr, font_name=cfg.body_font, font_size=FONT_SIZE_WU)

        # 页脚：水印文字 + 页码
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 水印文字
        fr_watermark = fp.add_run(cfg.footer_text)
        _set_run_font(fr_watermark, font_name=cfg.body_font, font_size=Pt(9),
                     color=RGBColor(0xCC, 0xCC, 0xCC))

        # 页码
        fp2 = footer.add_paragraph()
        fp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_page_number(fp2, cfg)


def _add_page_number(paragraph: docx.text.paragraph.Paragraph, cfg: WordFormatConfig) -> None:
    """在段落中插入 PAGE 域（阿拉伯数字页码，底部居中）"""
    run = paragraph.add_run()
    _set_run_font(run, font_name=cfg.body_font, font_size=FONT_SIZE_WU)

    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    run._r.append(fld_begin)

    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' PAGE '
    run._r.append(instr)

    fld_separate = OxmlElement('w:fldChar')
    fld_separate.set(qn('w:fldCharType'), 'separate')
    run._r.append(fld_separate)

    r2 = paragraph.add_run("1")
    _set_run_font(r2, font_name=cfg.body_font, font_size=FONT_SIZE_WU)

    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    r2._r.append(fld_end)


def _add_toc_field(doc: Document, cfg: WordFormatConfig) -> None:
    """插入目录域（TOC）"""
    # 目录标题
    toc_heading = doc.add_paragraph()
    toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = toc_heading.add_run(cfg.toc_title)
    _set_run_font(run, font_name=cfg.h1_font, font_size=FONT_SIZE_XIAO_ER, bold=True)
    toc_heading.paragraph_format.space_after = Pt(12)

    # 插入 TOC 域
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    run._r.append(fld_begin)

    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = f' TOC \\o "1-{cfg.toc_levels}" \\h \\z \\u '
    run._r.append(instr)

    fld_separate = OxmlElement('w:fldChar')
    fld_separate.set(qn('w:fldCharType'), 'separate')
    run._r.append(fld_separate)

    run2 = paragraph.add_run("（打开 Word 后按 Ctrl+A → F9 更新目录）")
    _set_run_font(run2, font_name=cfg.body_font, font_size=FONT_SIZE_WU,
                 color=RGBColor(0x99, 0x99, 0x99))

    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run2._r.append(fld_end)

    doc.add_page_break()


def _add_heading(doc: Document, text: str, level: int, cfg: WordFormatConfig) -> None:
    """添加带格式的标题"""
    heading = doc.add_heading(text, level=level)

    if level == 1:
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        font_name = cfg.h1_font
        font_size = Pt(cfg.h1_size_pt)
    elif level == 2:
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        font_name = cfg.h2_font
        font_size = Pt(cfg.h2_size_pt)
    else:
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        font_name = cfg.h3_font
        font_size = Pt(cfg.h3_size_pt)

    for run in heading.runs:
        _set_run_font(run, font_name=font_name, font_size=font_size, bold=True)

    # 段前段后间距
    pf = heading.paragraph_format
    if level == 1:
        pf.space_before = Pt(cfg.h1_before)
        pf.space_after = Pt(cfg.h1_after)
    elif level == 2:
        pf.space_before = Pt(cfg.h2_before)
        pf.space_after = Pt(cfg.h2_after)
    else:
        pf.space_before = Pt(cfg.h3_before)
        pf.space_after = Pt(cfg.h3_after)


def _add_paragraph(
    doc: Document,
    text: str,
    cfg: WordFormatConfig,
    first_indent: bool = True,
    bold: bool = False,
    alignment: Optional[int] = None,
) -> None:
    """添加正文段落"""
    if not text.strip():
        return
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    run = p.add_run(text.strip())
    _set_run_font(run, font_name=cfg.body_font, font_size=Pt(cfg.body_size_pt), bold=bold)
    _apply_paragraph_format(p, cfg, first_indent=first_indent)


def _add_table_from_markdown(
    doc: Document,
    markdown_table_lines: List[str],
    cfg: WordFormatConfig,
    table_fmt: TableFormatConfig,
) -> None:
    """将 Markdown 表格转换为 Word 表格"""
    if len(markdown_table_lines) < 2:
        return

    # 解析表头
    header_row = [c.strip() for c in markdown_table_lines[0].strip("|").split("|")]

    # 跳过分隔行（如 |---|---|），收集数据行
    data_rows = []
    for line in markdown_table_lines[1:]:
        line = line.strip()
        if not line or all(c in "-| :" for c in line):
            continue
        data_rows.append([c.strip() for c in line.strip("|").split("|")])

    if not header_row:
        return

    all_rows = [header_row] + data_rows
    num_cols = len(header_row)

    table = doc.add_table(rows=len(all_rows), cols=num_cols)
    table.style = "Table Grid"
    table.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER

    for row_idx, row_data in enumerate(all_rows):
        row = table.rows[row_idx]
        for col_idx, cell_text in enumerate(row_data):
            if col_idx >= num_cols:
                break
            cell = row.cells[col_idx]
            # 清除默认段落
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(cell_text)
            _set_run_font(
                run,
                font_name=table_fmt.font_name,
                font_size=table_fmt.font_size,
                bold=(row_idx == 0),
                color=table_fmt.header_font_color if row_idx == 0 else None,
            )
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 设置背景色
            if row_idx == 0:
                _set_cell_bg(cell, table_fmt.header_bg_color)
            elif row_idx % 2 == 0:
                _set_cell_bg(cell, table_fmt.alt_row_bg_color)

    # 表后空行
    doc.add_paragraph()


def _set_cell_bg(cell, color: RGBColor) -> None:
    """设置单元格背景色"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), f'{color[0]:02X}{color[1]:02X}{color[2]:02X}')
    shading.set(qn('w:val'), 'clear')
    tcPr.append(shading)


def _render_mermaid_as_image(
    doc: Document,
    mermaid_code: str,
    cfg: WordFormatConfig,
) -> None:
    """尝试通过 Kroki API 渲染 Mermaid 为图片嵌入文档

    注意：此函数包含同步网络 IO（urllib.request.urlopen），
    由 async 调用方通过 asyncio.to_thread() 在线程池中执行，
    以避免阻塞事件循环。
    """
    import base64
    import zlib
    import urllib.request

    try:
        compressed = zlib.compress(mermaid_code.encode("utf-8"), 9)
        encoded = base64.urlsafe_b64encode(compressed).decode("utf-8")
        url = f"https://kroki.io/mermaid/svg/{encoded}"

        req = urllib.request.Request(url, headers={"User-Agent": "Yibiao/1.0"})
        with urllib.request.urlopen(req, timeout=10) as resp:
            svg_data = resp.read()
            status = resp.status

        if status == 200 and svg_data:
            # 保存 SVG 为临时文件，使用 PIL 转换
            tmp_svg = os.path.join(tempfile.gettempdir(), f"mermaid_{uuid.uuid4().hex[:8]}.svg")
            tmp_png = os.path.join(tempfile.gettempdir(), f"mermaid_{uuid.uuid4().hex[:8]}.png")
            with open(tmp_svg, "wb") as f:
                f.write(svg_data)

            try:
                from PIL import Image as PILImage
                import cairosvg
                cairosvg.svg2png(url=tmp_svg, write_to=tmp_png)
                img = PILImage.open(tmp_png)
                # 缩放
                max_w = 540
                orig_w, orig_h = img.size
                if orig_w > max_w:
                    ratio = max_w / orig_w
                    img = img.resize((int(orig_w * ratio), int(orig_h * ratio)))
                img_path = tmp_png
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run()
                r.add_picture(img_path, width=Inches(min(5.5, max_w / 96)))
            except ImportError:
                # cairosvg 不可用，回退到代码块
                _render_mermaid_fallback(doc, mermaid_code, cfg)
            finally:
                for f in (tmp_svg, tmp_png):
                    if os.path.exists(f):
                        try:
                            os.remove(f)
                        except OSError:
                            pass
            return

    except Exception as e:
        logger.warning(f"Kroki API 不可用，保留 Mermaid 代码块: {e}")

    # Kroki 不可用 → 降级为代码块文本
    _render_mermaid_fallback(doc, mermaid_code, cfg)


def _render_mermaid_fallback(
    doc: Document,
    mermaid_code: str,
    cfg: WordFormatConfig,
) -> None:
    """降级方案：将 Mermaid 代码以等宽字体段落形式插入"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    note_run = p.add_run("[图表：Mermaid 代码 - 请在编辑时替换为正式图表]")
    _set_run_font(note_run, font_name=cfg.body_font, font_size=Pt(9),
                 color=RGBColor(0x99, 0x99, 0x99))

    for line in mermaid_code.strip().split("\n"):
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r2 = p2.add_run(line.rstrip())
        _set_run_font(r2, font_name="Courier New", font_size=Pt(8),
                     color=RGBColor(0x66, 0x66, 0x66))
        p2.paragraph_format.left_indent = Cm(1)
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(0)
        p2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE


# Markdown 解析正则
HEADING_RE = re.compile(r"^(#{1,6})\s+(.+)$", re.MULTILINE)
TABLE_ROW_RE = re.compile(r"^\|.+\|$")
TABLE_SEP_RE = re.compile(r"^\|[\s\-:|]+\|$")
MERMAID_START_RE = re.compile(r"^```mermaid\s*$")
CODE_BLOCK_END_RE = re.compile(r"^```\s*$")


def _render_markdown_to_doc(
    doc: Document,
    markdown: str,
    cfg: WordFormatConfig,
    table_fmt: Optional[TableFormatConfig] = None,
) -> None:
    """将 Markdown 内容渲染到 Word 文档

    支持：标题、正文、表格、Mermaid 图表代码块、Material Suggestions
    """
    if not markdown:
        return

    tf = table_fmt or TableFormatConfig()
    lines = markdown.split("\n")
    i = 0

    while i < len(lines):
        line = lines[i]

        # 跳过空行
        if not line.strip():
            i += 1
            continue

        # Material Suggestions 块（HTML 注释标记）
        if "<!-- MATERIAL_SUGGESTIONS -->" in line:
            # 跳过整个 Material Suggestions 块（已作为 normal markdown 渲染为表格）
            i += 1
            continue

        # Mermaid 代码块
        if MERMAID_START_RE.match(line.strip()):
            mermaid_lines = []
            i += 1
            while i < len(lines) and not CODE_BLOCK_END_RE.match(lines[i].strip()):
                mermaid_lines.append(lines[i])
                i += 1
            if mermaid_lines:
                # _render_mermaid_as_image 含同步 IO，但在当前同步渲染上下文中
                # 整个 _render_markdown_to_doc 本身是由 export_to_docx 在
                # asyncio.to_thread 中调用的，因此这里直接执行是安全的。
                _render_mermaid_as_image(doc, "\n".join(mermaid_lines), cfg)
            i += 1
            continue

        # 标题
        heading_match = HEADING_RE.match(line)
        if heading_match:
            level = len(heading_match.group(1))
            title = heading_match.group(2).strip()
            # 跳过 Material Suggestions 二级标题
            if "建议引用素材" in title:
                i += 1
                # 跳过接下的 Markdown 表格行直到空行
                while i < len(lines) and lines[i].strip():
                    i += 1
                continue
            _add_heading(doc, title, min(level, 3), cfg)
            i += 1
            continue

        # Markdown 表格
        if TABLE_ROW_RE.match(line.strip()):
            table_lines = []
            while i < len(lines) and (
                TABLE_ROW_RE.match(lines[i].strip())
                or TABLE_SEP_RE.match(lines[i].strip())
            ):
                table_lines.append(lines[i])
                i += 1
            if table_lines:
                _add_table_from_markdown(doc, table_lines, cfg, tf)
            continue

        # 水平线
        if line.strip() in ("---", "***", "___"):
            doc.add_paragraph()
            i += 1
            continue

        # 普通段落
        # 去除内联 Markdown 标记（**bold**, *italic*）
        clean_text = re.sub(r"\*\*(.+?)\*\*", r"\1", line)
        clean_text = re.sub(r"\*(.+?)\*", r"\1", clean_text)
        clean_text = re.sub(r"`(.+?)`", r"\1", clean_text)

        _add_paragraph(doc, clean_text, cfg)
        i += 1


class WordExportService:
    """Word 文档导出服务

    从 Markdown 内容生成符合 GB/T 9704 格式规范的 Word 文档。
    默认使用标书标准样式：
    - # → Heading 1：宋体 小二 加粗 居中
    - ## → Heading 2：黑体 三号 加粗
    - ### → Heading 3：黑体 小三 加粗
    - 正文：宋体 小四 1.5 倍行距 首行缩进 2 字符
    - 页眉：项目名称居中
    - 页脚："严格保密" 水印 + 页码居中
    - 目录：自动 TOC 域
    """

    def __init__(self):
        pass

    @staticmethod
    async def load_format_config(
        db: AsyncSession,
        template_id: Optional[str] = None,
    ) -> WordFormatConfig:
        """从数据库加载格式配置"""
        if not template_id:
            return WordFormatConfig()

        try:
            tid = uuid.UUID(template_id)
        except ValueError:
            return WordFormatConfig()

        from ..models.export_template import ExportTemplate
        tmpl = await db.get(ExportTemplate, tid)
        if not tmpl:
            return WordFormatConfig()

        return _parse_word_config_from_template(tmpl.format_config)

    @staticmethod
    def _build_docx(
        project_name: str,
        chapters: List[Dict[str, Any]],
        project_overview: str,
        cfg: WordFormatConfig,
    ) -> io.BytesIO:
        """同步构建 Word 文档（含 Mermaid 渲染等同步 IO）

        此方法由 export_to_docx 通过 asyncio.to_thread 在线程池中调用，
        避免阻塞事件循环。不要直接在 async 上下文中调用此方法。
        """
        doc = Document()

        # 初始化 Normal 样式
        styles = doc.styles
        if "Normal" in styles:
            style = styles["Normal"]
            style.font.name = cfg.body_font
            style.font.size = Pt(cfg.body_size_pt)
            style.font.bold = False
            if style._element.rPr is None:
                style._element._add_rPr()
            style._element.rPr.rFonts.set(qn("w:eastAsia"), cfg.body_font)

        # 设置页面
        _setup_section(doc, cfg)

        # 页眉和页脚
        _setup_header_footer(doc, project_name, cfg)

        # 目录
        if cfg.show_toc:
            _add_toc_field(doc, cfg)

        # 项目概述
        if project_overview:
            _add_heading(doc, "项目概述", 1, cfg)
            _add_paragraph(doc, project_overview, cfg)

        # 渲染各章节
        table_fmt = TableFormatConfig(
            font_name=cfg.table_font,
            font_size=Pt(cfg.table_size_pt),
        )

        for ch in chapters:
            title = ch.get("title", "")
            content = ch.get("content", "") or ch.get("text", "") or ""

            if title:
                _add_heading(doc, title, 1, cfg)
            if content:
                _render_markdown_to_doc(doc, content, cfg, table_fmt)

        # 保存到缓冲区
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    async def export_to_docx(
        self,
        project_name: str,
        chapters: List[Dict[str, Any]],
        project_overview: str = "",
        format_config: Optional[WordFormatConfig] = None,
    ) -> io.BytesIO:
        """将章节内容导出为 Word 文档

        Args:
            project_name: 项目名称
            chapters: 章节列表 [{title, content}, ...]
            project_overview: 项目概述
            format_config: 格式配置（可选，默认使用标书标准格式）

        Returns:
            包含 docx 数据的 BytesIO
        """
        import asyncio

        cfg = format_config or WordFormatConfig()

        # 将同步的文档构建（含 urllib 网络调用）放到线程池执行
        return await asyncio.to_thread(
            self._build_docx,
            project_name,
            chapters,
            project_overview,
            cfg,
        )
