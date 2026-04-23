"""素材库服务"""
from __future__ import annotations

import os
import re
import uuid
from datetime import date, datetime, timezone
from io import BytesIO
from typing import Iterable

from docx.document import Document as DocxDocument
from docx.shared import Inches

from ..models.material import BindingDisplayMode, MaterialAsset
from ..schemas.material import MaterialAssetResponse, MaterialRequirementMatchCandidate

MATERIAL_PATTERN = re.compile(r"\[INSERT_MATERIAL:([a-zA-Z0-9-]+)\]")


def build_material_storage_paths(*, scope: str, owner_id: uuid.UUID, material_id: uuid.UUID, extension: str) -> dict[str, str]:
    base_dir = os.path.join(
        "uploads",
        "materials",
        scope,
        str(owner_id),
        str(material_id),
    )
    normalized_ext = extension.lower().lstrip(".")
    return {
        "base_dir": base_dir,
        "original": os.path.join(base_dir, f"original.{normalized_ext}"),
        "preview": os.path.join(base_dir, "preview.jpg"),
        "thumbnail": os.path.join(base_dir, "thumb.jpg"),
    }


def _normalize_tokens(values: Iterable[str] | None) -> set[str]:
    tokens: set[str] = set()
    if not values:
        return tokens
    for value in values:
        normalized = value.strip().lower()
        if normalized:
            tokens.add(normalized)
    return tokens


def _recency_bonus(last_used_at: datetime | None, *, now: datetime) -> int:
    if not last_used_at:
        return 0
    delta_days = max((now - last_used_at).days, 0)
    if delta_days <= 30:
        return 5
    if delta_days <= 90:
        return 3
    return 1


def _asset_to_response(asset: MaterialAsset) -> MaterialAssetResponse:
    created_at = asset.created_at or datetime.now(timezone.utc)
    updated_at = asset.updated_at or created_at
    return MaterialAssetResponse(
        id=str(asset.id),
        scope=asset.scope,
        owner_id=str(asset.owner_id) if asset.owner_id else None,
        uploaded_by=str(asset.uploaded_by) if asset.uploaded_by else None,
        category=asset.category,
        name=asset.name,
        description=asset.description,
        file_path=asset.file_path,
        preview_path=asset.preview_path,
        thumbnail_path=asset.thumbnail_path,
        file_type=asset.file_type,
        file_ext=asset.file_ext,
        file_size=asset.file_size,
        page_count=asset.page_count,
        tags=list(asset.tags or []),
        keywords=list(asset.keywords or []),
        ai_description=asset.ai_description,
        ai_extracted_fields=dict(asset.ai_extracted_fields or {}),
        valid_from=asset.valid_from,
        valid_until=asset.valid_until,
        is_expired=asset.is_expired,
        is_disabled=asset.is_disabled or False,
        review_status=asset.review_status,
        usage_count=asset.usage_count,
        last_used_at=asset.last_used_at,
        created_at=created_at,
        updated_at=updated_at,
    )


def build_material_requirement_candidates(
    *,
    requirement_name: str,
    requirement_text: str,
    requirement_category: str | None,
    requirement_tags: list[str] | None,
    assets: list[MaterialAsset],
    today: date | None = None,
) -> list[MaterialRequirementMatchCandidate]:
    current_day = today or date.today()
    now = datetime.now(timezone.utc)
    required_text = f"{requirement_name} {requirement_text}".lower()
    requirement_tag_set = _normalize_tokens(requirement_tags)
    candidates: list[MaterialRequirementMatchCandidate] = []

    for asset in assets:
        if asset.valid_until and asset.valid_until < current_day:
            continue

        score = 0.0
        reasons: list[str] = []

        if requirement_category and asset.category.value == requirement_category:
            score += 40
            reasons.append("分类完全匹配")

        asset_tag_set = _normalize_tokens(asset.tags)
        matched_tags = requirement_tag_set & asset_tag_set
        if matched_tags:
            score += min(len(matched_tags) * 15, 30)
            reasons.append(f"标签命中: {', '.join(sorted(matched_tags))}")

        company_name = str((asset.ai_extracted_fields or {}).get("company_name", "")).strip()
        if company_name and company_name.lower() in required_text:
            score += 20
            reasons.append("主体名称命中")

        if asset.name.lower() in required_text or any(token in asset.name.lower() for token in requirement_tag_set):
            score += 15
            reasons.append("素材名称相关")

        score += _recency_bonus(asset.last_used_at, now=now)

        candidates.append(
                MaterialRequirementMatchCandidate(
                    asset_id=str(asset.id),
                    score=score,
                    matched_reasons=reasons,
                    asset=_asset_to_response(asset),
                )
            )

    return sorted(candidates, key=lambda item: item.score, reverse=True)[:5]


def render_material_block(document: DocxDocument, binding: dict) -> None:
    display_mode = binding.get("display_mode") or BindingDisplayMode.IMAGE.value
    caption = binding.get("caption") or binding.get("material_name") or "素材附件"
    image_path = binding.get("preview_path") or binding.get("file_path")
    file_type = (binding.get("file_type") or "").lower()

    if display_mode == BindingDisplayMode.ATTACHMENT_NOTE.value:
        document.add_paragraph(f"[附件说明] {caption}")
        return

    if image_path and os.path.exists(image_path):
        # 按宽高比自适应缩放
        max_width = Inches(5.5)
        max_height = Inches(8.0)
        min_width = Inches(2.0)
        img_width = max_width  # 默认值

        try:
            from PIL import Image as PILImage
            with PILImage.open(image_path) as img:
                orig_w, orig_h = img.size  # 像素
                if orig_w > 0 and orig_h > 0:
                    aspect = orig_h / orig_w
                    # 先按最大宽度计算
                    target_width = max_width
                    target_height = Inches(max_width.inches * aspect)
                    # 如果高度超限，按最大高度反推宽度
                    if target_height > max_height:
                        target_width = Inches(max_height.inches / aspect)
                        target_height = max_height
                    # 如果原图比最大值小，用原图尺寸（DPI 按 96 估算）
                    orig_w_inches = orig_w / 96.0
                    if orig_w_inches < max_width.inches:
                        target_width = Inches(max(orig_w_inches, min_width.inches))
                        target_height = Inches(target_width.inches * aspect)
                        if target_height > max_height:
                            target_width = Inches(max_height.inches / aspect)
                    img_width = target_width
        except Exception:
            pass  # PIL 不可用时使用默认宽度

        paragraph = document.add_paragraph()
        run = paragraph.add_run()
        run.add_picture(image_path, width=img_width)
        if "pdf" in file_type:
            document.add_paragraph(f"{caption}（PDF 附件预览）")
        else:
            document.add_paragraph(caption)
        return

    missing_name = binding.get("material_name") or caption
    document.add_paragraph(f"[素材缺失：{missing_name}]")


def render_content_with_material_bindings(
    *,
    document: DocxDocument,
    content: str,
    binding_map: dict[str, dict],
    add_text: callable | None = None,
) -> None:
    cursor = 0

    def append_text(text: str) -> None:
        cleaned = text.strip()
        if not cleaned:
            return
        if add_text:
            add_text(cleaned)
        else:
            for line in cleaned.split("\n"):
                if line.strip():
                    document.add_paragraph(line.strip())

    for match in MATERIAL_PATTERN.finditer(content):
        append_text(content[cursor:match.start()])
        binding_id = match.group(1)
        binding = binding_map.get(binding_id)
        if binding:
            render_material_block(document, binding)
        else:
            document.add_paragraph("[素材缺失：未找到绑定配置]")
        cursor = match.end()

    append_text(content[cursor:])
