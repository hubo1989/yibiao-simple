"""历史标书解析服务"""
from __future__ import annotations

import asyncio
import logging
import os
import re
import uuid
from datetime import datetime, timezone
from typing import TYPE_CHECKING

from sqlalchemy import select, update
from sqlalchemy.ext.asyncio import AsyncSession

from ..models.ingestion import IngestionTask, IngestionTaskStatus, MaterialCandidate
from ..models.knowledge import KnowledgeDoc
from ..models.material import MaterialAsset, MaterialCategory, MaterialReviewStatus

if TYPE_CHECKING:
    from pathlib import Path

logger = logging.getLogger(__name__)

# 规则匹配模式
RULE_PATTERNS = {
    MaterialCategory.BUSINESS_LICENSE: ["营业执照", "统一社会信用代码", "工商登记"],
    MaterialCategory.QUALIFICATION_CERT: ["资质证书", "资质等级", "建筑业企业资质", "设计资质", "施工资质"],
    MaterialCategory.ISO_CERT: ["ISO", "质量管理体系", "环境管理体系", "职业健康安全"],
    MaterialCategory.CONTRACT_SAMPLE: ["合同", "甲方", "乙方", "签订日期", "合同金额"],
    MaterialCategory.PROJECT_CASE: ["项目案例", "业绩", "成功案例", "工程业绩", "施工业绩"],
    MaterialCategory.AWARD_CERT: ["获奖证书", "荣誉证书", "优质工程", "鲁班奖"],
    MaterialCategory.TEAM_PHOTO: ["团队", "人员", "项目经理", "技术负责人"],
    MaterialCategory.EQUIPMENT_PHOTO: ["设备", "机械", "仪器", "车辆"],
}

# 结构化抽取 schema
EXTRACTION_SCHEMAS = {
    MaterialCategory.BUSINESS_LICENSE: {
        "company_name": "公司名称",
        "credit_code": "统一社会信用代码",
        "legal_person": "法定代表人",
        "registered_capital": "注册资本",
        "valid_until": "有效期至",
    },
    MaterialCategory.CONTRACT_SAMPLE: {
        "project_name": "项目名称",
        "party_a": "甲方",
        "party_b": "乙方",
        "sign_date": "签订日期",
        "contract_amount": "合同金额",
        "project_location": "项目地点",
    },
    MaterialCategory.PROJECT_CASE: {
        "project_name": "项目名称",
        "project_owner": "业主单位",
        "project_scale": "项目规模",
        "contract_amount": "合同金额",
        "start_date": "开工日期",
        "completion_date": "竣工日期",
        "project_manager": "项目经理",
    },
    MaterialCategory.QUALIFICATION_CERT: {
        "company_name": "企业名称",
        "certificate_number": "证书编号",
        "qualification_grade": "资质等级",
        "valid_until": "有效期至",
    },
}


class HistoricalBidIngestionService:
    """历史标书解析服务"""

    def __init__(self, db: AsyncSession):
        self.db = db

    async def create_ingestion_task(self, document_id: uuid.UUID, created_by: uuid.UUID | None = None) -> IngestionTask:
        """创建解析任务"""
        # 检查文档是否存在
        doc = await self.db.get(KnowledgeDoc, document_id)
        if not doc:
            raise ValueError(f"文档不存在: {document_id}")

        # 检查是否已有进行中的任务——超过 5 分钟的视为僵死，自动清理
        from datetime import timedelta
        stale_threshold = datetime.now(timezone.utc) - timedelta(minutes=5)
        existing_result = await self.db.execute(
            select(IngestionTask).where(
                IngestionTask.document_id == document_id,
                IngestionTask.status.in_([IngestionTaskStatus.PENDING, IngestionTaskStatus.PROCESSING]),
            )
        )
        existing_task = existing_result.scalar_one_or_none()
        if existing_task:
            if existing_task.created_at < stale_threshold:
                # 僵死任务，自动标记为失败
                existing_task.status = IngestionTaskStatus.FAILED
                existing_task.error_message = "任务超时，已自动清理"
                await self.db.commit()
                logger.warning(f"自动清理僵死任务: {existing_task.id}")
            else:
                raise ValueError("该文档已有进行中的解析任务，请稍后再试")

        task = IngestionTask(
            document_id=document_id,
            created_by=created_by,
            status=IngestionTaskStatus.PENDING,
        )
        self.db.add(task)
        await self.db.commit()
        await self.db.refresh(task)

        logger.info(f"创建解析任务: {task.id}, 文档: {document_id}")
        return task

    async def process_document(self, task_id: uuid.UUID) -> list[MaterialCandidate]:
        """处理文档并生成候选素材"""
        task = await self.db.get(IngestionTask, task_id)
        if not task:
            raise ValueError(f"任务不存在: {task_id}")

        # 更新任务状态
        task.status = IngestionTaskStatus.PROCESSING
        task.started_at = datetime.now(timezone.utc)
        task.processing_log = task.processing_log or []
        task.processing_log.append({
            "time": datetime.now(timezone.utc).isoformat(),
            "action": "start_processing",
            "message": "开始解析文档",
        })
        await self.db.commit()

        try:
            # 获取文档
            doc = await self.db.get(KnowledgeDoc, task.document_id)
            if not doc:
                raise ValueError(f"文档不存在: {task.document_id}")

            # 解析文档内容
            candidates = await self._extract_candidates(doc, task)

            # 保存候选素材
            for candidate in candidates:
                self.db.add(candidate)

            task.total_candidates = len(candidates)
            task.status = IngestionTaskStatus.COMPLETED
            task.processing_log.append({
                "time": datetime.now(timezone.utc).isoformat(),
                "action": "extraction_complete",
                "message": f"提取完成，共 {len(candidates)} 个候选素材",
            })

            await self.db.commit()
            logger.info(f"任务 {task_id} 解析完成，提取 {len(candidates)} 个候选素材")

            return candidates

        except Exception as e:
            task.status = IngestionTaskStatus.FAILED
            task.error_message = str(e)
            task.processing_log.append({
                "time": datetime.now(timezone.utc).isoformat(),
                "action": "error",
                "message": str(e),
            })
            await self.db.commit()
            logger.error(f"任务 {task_id} 解析失败: {e}")
            raise

    async def _extract_candidates(self, doc: KnowledgeDoc, task: IngestionTask) -> list[MaterialCandidate]:
        """从文档中提取候选素材"""
        candidates = []

        # 1. 读取文档文件
        file_path = doc.file_path
        if not os.path.exists(file_path):
            logger.warning(f"文档文件不存在: {file_path}")
            return candidates

        # 2. 根据文件类型解析
        if file_path.endswith(".pdf"):
            page_contents = await self._parse_pdf(file_path)
        elif file_path.endswith((".docx", ".doc")):
            page_contents = await self._parse_docx(file_path)
        else:
            logger.warning(f"不支持的文件类型: {file_path}")
            return candidates

        # 3. 规则召回
        for page_num, (text, images) in enumerate(page_contents, start=1):
            # 规则匹配
            matched_categories = self._rule_match(text)

            for category, confidence in matched_categories.items():
                # 创建候选素材
                candidate = MaterialCandidate(
                    task_id=task.id,
                    category=category.value,
                    name=self._generate_name(category, text, page_num),
                    source_page_from=page_num,
                    source_page_to=page_num,
                    source_excerpt=text[:500] if text else None,
                    extraction_method="rule",
                    confidence_score=confidence,
                    ai_extracted_fields=self._extract_fields(category, text),
                    review_status="pending",
                )
                candidates.append(candidate)

        # 4. LLM 增强（可选）
        # candidates = await self._llm_classify(candidates, page_contents)

        return candidates

    async def _parse_pdf(self, file_path: str) -> list[tuple[str, list]]:
        """解析 PDF 文件，返回每页的 (文本, 图片列表)"""
        import asyncio
        return await asyncio.to_thread(self._parse_pdf_sync, file_path)

    def _parse_pdf_sync(self, file_path: str) -> list[tuple[str, list]]:
        """同步解析 PDF"""
        try:
            import fitz  # PyMuPDF
        except ImportError:
            logger.warning("PyMuPDF 未安装，无法解析 PDF")
            return []

        page_contents = []
        doc = fitz.open(file_path)

        for page_num in range(len(doc)):
            page = doc[page_num]
            text = page.get_text()

            images = []
            for img in page.get_images():
                xref = img[0]
                base_image = doc.extract_image(xref)
                images.append({
                    "xref": xref,
                    "ext": base_image["ext"],
                    "data": base_image["image"],
                })

            page_contents.append((text, images))

        doc.close()
        return page_contents

    async def _parse_docx(self, file_path: str) -> list[tuple[str, list]]:
        """解析 DOCX 文件，按大纲级别分段，提取图片"""
        import asyncio
        return await asyncio.to_thread(self._parse_docx_sync, file_path)

    def _parse_docx_sync(self, file_path: str) -> list[tuple[str, list]]:
        """同步解析 DOCX，按标题拆分为逻辑分段"""
        from docx import Document
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        import zipfile
        import os

        doc = Document(file_path)

        # 提取所有图片（通过 rels）
        image_map: dict[str, dict] = {}
        try:
            with zipfile.ZipFile(file_path, 'r') as zf:
                for rel in doc.part.rels.values():
                    if "image" in rel.reltype:
                        target = rel.target_ref
                        # target 可能是 "media/image1.png" 或 "word/media/image1.png"
                        for prefix in [f"word/{target}", target]:
                            if prefix in zf.namelist():
                                data = zf.read(prefix)
                                ext = os.path.splitext(target)[1].lstrip('.')
                                image_map[rel.rId] = {
                                    "ext": ext or "png",
                                    "data": data,
                                    "xref": rel.rId,
                                }
                                break
        except Exception as e:
            logger.warning(f"DOCX 图片提取失败: {e}")

        # 按标题分段
        sections: list[tuple[str, list]] = []
        current_lines: list[str] = []
        current_images: list[dict] = []

        def _flush_section():
            nonlocal current_lines, current_images
            text = "\n".join(current_lines).strip()
            if text:
                sections.append((text, current_images))
            current_lines = []
            current_images = []

        for para in doc.paragraphs:
            # 检查是否是标题段落（Heading 样式）
            style_name = para.style.name if para.style else ""
            is_heading = style_name.startswith("Heading") or style_name.startswith("标题")

            if is_heading and current_lines:
                _flush_section()

            if para.text.strip():
                current_lines.append(para.text)

            # 检查段落中的图片引用
            for run in para.runs:
                for drawing in run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'):
                    # 从 drawing 中提取 blip rId
                    blips = drawing.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
                    for blip in blips:
                        embed = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                        if embed and embed in image_map:
                            current_images.append(image_map[embed])

        # 最后一段
        _flush_section()

        # 如果没有标题分段，退化为单段
        if not sections:
            all_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            all_imgs = list(image_map.values())
            if all_text:
                sections = [(all_text, all_imgs)]

        return sections

    def _rule_match(self, text: str) -> dict[MaterialCategory, float]:
        """规则匹配，返回类别和置信度"""
        matches = {}
        text_lower = text.lower()

        for category, keywords in RULE_PATTERNS.items():
            score = sum(1 for kw in keywords if kw in text_lower)
            if score > 0:
                confidence = min(score / len(keywords) + 0.3, 0.95)
                matches[category] = confidence

        return matches

    def _generate_name(self, category: MaterialCategory, text: str, page_num: int) -> str:
        """生成候选素材名称"""
        category_names = {
            MaterialCategory.BUSINESS_LICENSE: "营业执照",
            MaterialCategory.QUALIFICATION_CERT: "资质证书",
            MaterialCategory.ISO_CERT: "ISO认证证书",
            MaterialCategory.CONTRACT_SAMPLE: "合同",
            MaterialCategory.PROJECT_CASE: "项目案例",
            MaterialCategory.AWARD_CERT: "获奖证书",
            MaterialCategory.TEAM_PHOTO: "团队照片",
            MaterialCategory.EQUIPMENT_PHOTO: "设备照片",
        }
        base_name = category_names.get(category, "素材")

        # 尝试从文本中提取关键信息
        if category == MaterialCategory.PROJECT_CASE:
            # 提取项目名称
            match = re.search(r"项目名称[：:]\s*([^\n]+)", text)
            if match:
                return f"{match.group(1).strip()[:50]}"
        elif category == MaterialCategory.CONTRACT_SAMPLE:
            # 提取合同名称
            match = re.search(r"合同名称[：:]\s*([^\n]+)", text)
            if match:
                return f"{match.group(1).strip()[:50]}"

        return f"{base_name} - 第{page_num}页"

    def _extract_fields(self, category: MaterialCategory, text: str) -> dict:
        """结构化字段抽取"""
        schema = EXTRACTION_SCHEMAS.get(category, {})
        fields = {}

        for field_name, label in schema.items():
            # 尝试多种模式匹配
            patterns = [
                rf"{label}[：:]\s*([^\n]+)",
                rf"{label}[:：]\s*([^\n]+)",
                rf"{label}\s*[:：]\s*([^\n]+)",
            ]

            for pattern in patterns:
                match = re.search(pattern, text)
                if match:
                    value = match.group(1).strip()
                    # 清理常见噪音
                    value = re.sub(r"\s+", " ", value)
                    if len(value) > 0 and len(value) < 200:
                        fields[field_name] = value
                    break

        return fields if fields else None

    async def confirm_candidates(
        self,
        task_id: uuid.UUID,
        confirm_ids: list[uuid.UUID],
        reject_ids: list[uuid.UUID],
        owner_id: uuid.UUID | None = None,
        scope: str = "user",
    ) -> tuple[int, int]:
        """确认候选素材并入库"""
        task = await self.db.get(IngestionTask, task_id)
        if not task:
            raise ValueError(f"任务不存在: {task_id}")

        confirmed_count = 0
        rejected_count = 0

        # 确认入库
        for candidate_id in confirm_ids:
            candidate = await self.db.get(MaterialCandidate, candidate_id)
            if not candidate or candidate.task_id != task_id:
                continue

            # 创建正式素材
            asset = MaterialAsset(
                scope=scope,
                owner_id=owner_id,
                category=candidate.category,
                name=candidate.name,
                source_document_id=task.document_id,
                source_page_from=candidate.source_page_from,
                source_page_to=candidate.source_page_to,
                source_excerpt=candidate.source_excerpt,
                extraction_method=candidate.extraction_method,
                ai_description=candidate.ai_description,
                ai_extracted_fields=candidate.ai_extracted_fields,
                tags=candidate.tags or [],
                review_status=MaterialReviewStatus.PENDING,
                file_path=candidate.temp_file_path or "",
                file_type=candidate.file_type or "application/octet-stream",
                file_ext=candidate.file_ext or "",
                file_size=candidate.file_size or 0,
            )

            # 尝试映射 category 到 enum
            try:
                asset.category = MaterialCategory(candidate.category)
            except ValueError:
                asset.category = MaterialCategory.OTHER

            self.db.add(asset)

            # 标记候选为已确认
            candidate.review_status = "confirmed"
            candidate.confirmed_at = datetime.now(timezone.utc)
            candidate.confirmed_by = owner_id

            confirmed_count += 1

        # 拒绝
        for candidate_id in reject_ids:
            candidate = await self.db.get(MaterialCandidate, candidate_id)
            if not candidate or candidate.task_id != task_id:
                continue

            candidate.review_status = "rejected"
            rejected_count += 1

        # 更新任务统计
        task.confirmed_count = (task.confirmed_count or 0) + confirmed_count
        task.rejected_count = (task.rejected_count or 0) + rejected_count

        # 检查是否全部处理完成
        if task.confirmed_count + task.rejected_count >= task.total_candidates:
            task.status = IngestionTaskStatus.COMPLETED
            task.completed_at = datetime.now(timezone.utc)

        await self.db.commit()
        logger.info(f"任务 {task_id}: 确认 {confirmed_count} 个，拒绝 {rejected_count} 个")

        return confirmed_count, rejected_count

    async def get_task_status(self, task_id: uuid.UUID) -> IngestionTask | None:
        """获取任务状态"""
        return await self.db.get(IngestionTask, task_id)

    async def list_candidates(self, task_id: uuid.UUID) -> list[MaterialCandidate]:
        """获取任务的所有候选素材"""
        result = await self.db.execute(
            select(MaterialCandidate).where(MaterialCandidate.task_id == task_id)
        )
        return list(result.scalars().all())
