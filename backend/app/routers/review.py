"""标书审查 API 路由"""
import os
import uuid
from datetime import datetime
from typing import Annotated
from urllib.parse import quote

from docx import Document
from fastapi import (
    APIRouter,
    Depends,
    File,
    Form,
    HTTPException,
    UploadFile,
    status,
)
from fastapi.responses import StreamingResponse
from sqlalchemy import and_, exists, select
from sqlalchemy.ext.asyncio import AsyncSession

from ..auth.dependencies import require_editor
from ..config import settings
from ..db.database import get_db
from ..models.bid_review_task import BidReviewTask, ReviewTaskStatus
from ..models.project import Project, project_members
from ..models.user import User
from ..schemas.review import (
    BidFileUploadResponse,
    ReviewExecuteRequest,
    ReviewExportRequest,
    ReviewHistoryResponse,
    ReviewResultResponse,
)
from ..services.file_service import FileService
from ..services.review_service import ReviewService
from ..services.word_comment_service import WordCommentService
from ..utils.sse import sse_response

import json

router = APIRouter(prefix="/api/review", tags=["标书审查"])


async def _verify_project_member(
    project_id: uuid.UUID,
    user_id: uuid.UUID,
    db: AsyncSession,
) -> Project:
    """验证用户是项目成员并返回项目"""
    member_exists = await db.execute(
        select(
            exists().where(
                and_(
                    project_members.c.project_id == project_id,
                    project_members.c.user_id == user_id,
                )
            )
        )
    )
    if not member_exists.scalar():
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="项目不存在或您没有访问权限",
        )

    result = await db.execute(select(Project).where(Project.id == project_id))
    project = result.scalar_one_or_none()
    if not project:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="项目不存在",
        )
    return project


def _build_paragraph_index(doc: Document) -> list[dict]:
    """从 python-docx Document 构建段落索引"""
    index = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue

        para_type = "paragraph"
        level = None
        style_name = para.style.name if para.style else ""

        if style_name.startswith("Heading"):
            para_type = "heading"
            try:
                level = int(style_name.split()[-1])
            except (ValueError, IndexError):
                pass

        entry: dict = {"index": i, "type": para_type, "text": text}
        if level is not None:
            entry["level"] = level
        index.append(entry)

    # 表格
    for i, table in enumerate(doc.tables):
        rows_text = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            row_text = " | ".join(cells)
            if row_text:
                rows_text.append(row_text)
        if rows_text:
            index.append({
                "index": f"table_{i}",
                "type": "table",
                "text": "\n".join(rows_text),
            })

    return index


# ============ 投标文件上传 ============


@router.post("/upload-bid", response_model=BidFileUploadResponse)
async def upload_bid_file(
    project_id: Annotated[str, Form(...)],
    file: UploadFile = File(...),
    current_user: Annotated[User, Depends(require_editor)] = None,
    db: Annotated[AsyncSession, Depends(get_db)] = None,
):
    """上传投标文件到指定项目的审查任务。

    接收 .docx 格式的投标文件，校验项目状态、文件类型和大小后，
    提取文本内容并构建段落索引，创建审查任务记录。

    前置条件：项目已上传招标文件且已完成分析。
    """
    try:
        project_uuid = uuid.UUID(project_id)
    except ValueError:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="无效的项目ID格式",
        )

    project = await _verify_project_member(project_uuid, current_user.id, db)

    # 前置条件校验：项目已上传招标文件且已完成分析
    if not project.file_content:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="项目尚未上传招标文件，请先上传",
        )
    if not project.project_overview or not project.tech_requirements:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="项目尚未完成招标文件分析，请先分析",
        )

    # 文件类型校验：仅支持 .docx
    docx_content_type = (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    if file.content_type != docx_content_type:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="仅支持 .docx 格式的投标文件",
        )

    # Magic bytes 校验
    header = await file.read(8)
    await file.seek(0)
    if header[:4] != b"PK\x03\x04":
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="文件内容与类型不匹配，请上传有效的 .docx 文件",
        )

    # 文件大小校验（50MB）
    max_size = 50 * 1024 * 1024
    file_content_bytes = await file.read()
    file_size = len(file_content_bytes)
    if file_size > max_size:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="文件大小超过 50MB 限制",
        )
    await file.seek(0)

    # 保存原始文件到磁盘
    review_dir = os.path.join(settings.upload_dir, "review", str(project_uuid))
    os.makedirs(review_dir, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
    # 安全处理客户端文件名：移除路径分隔符，使用原始文件名（不含扩展名）+ 安全扩展名
    original_name = os.path.basename(file.filename or "bid").split('.')[0]
    # 只保留安全字符（字母、数字、下划线、连字符、中文）
    safe_name = "".join(c for c in original_name if c.isalnum() or c in ('_', '-', ' ', '.', '，', '、'))
    safe_name = safe_name.strip()[:50]  # 限制长度
    safe_filename = f"{timestamp}_{safe_name}.docx"
    bid_file_path = os.path.join(review_dir, safe_filename)

    with open(bid_file_path, "wb") as f:
        f.write(file_content_bytes)

    # 提取文本内容
    bid_content = await FileService.extract_text_from_docx(bid_file_path)

    # 构建段落索引
    paragraph_index = []
    heading_count = 0
    try:
        doc = Document(bid_file_path)
        paragraph_index = _build_paragraph_index(doc)
        heading_count = sum(1 for p in paragraph_index if p.get("type") == "heading")
    except Exception as e:
        print(f"[Review] 构建段落索引失败: {e}")

    # 创建审查任务记录
    task = BidReviewTask(
        project_id=project_uuid,
        created_by=current_user.id,
        bid_file_path=bid_file_path,
        bid_filename=file.filename or "bid.docx",
        bid_content=bid_content,
        paragraph_index=paragraph_index,
        dimensions=["responsiveness", "compliance", "consistency"],
    )
    db.add(task)
    await db.flush()
    await db.refresh(task)

    return BidFileUploadResponse(
        success=True,
        message="投标文件上传成功",
        task_id=str(task.id),
        file_info={
            "filename": file.filename,
            "file_size": file_size,
            "paragraph_count": len(paragraph_index),
            "heading_count": heading_count,
        },
    )


# ============ 执行审查（SSE） ============


@router.post("/execute")
async def execute_review(
    request: ReviewExecuteRequest,
    current_user: Annotated[User, Depends(require_editor)] = None,
    db: Annotated[AsyncSession, Depends(get_db)] = None,
):
    """执行审查任务，SSE 流式返回进度和结果。

    支持响应性、合规性、一致性三个维度的并行审查。
    通过 SSE 事件流实时推送每个维度的执行进度和最终结果。
    """
    task = await db.get(BidReviewTask, request.task_id)
    if not task:
        raise HTTPException(status_code=404, detail="审查任务不存在")

    await _verify_project_member(task.project_id, current_user.id, db)

    if task.status == ReviewTaskStatus.PROCESSING:
        raise HTTPException(status_code=400, detail="审查任务正在执行中")

    # 更新审查配置
    task.dimensions = request.dimensions
    task.scope = request.scope
    task.model_name = request.model_name
    task.provider_config_id = request.provider_config_id
    await db.flush()

    task_id = request.task_id
    dimensions = request.dimensions

    async def generate():
        try:
            # 发送开始信号
            yield f"data: {json.dumps({'type': 'started', 'message': '开始执行审查...'}, ensure_ascii=False)}\n\n"

            # 为每个维度发送进度
            dimension_labels = {
                "responsiveness": "响应性",
                "compliance": "合规性",
                "consistency": "一致性",
            }

            for dim in dimensions:
                label = dimension_labels.get(dim, dim)
                yield f"data: {json.dumps({'type': 'progress', 'dimension': dim, 'status': 'processing', 'message': f'正在执行{label}审查...'}, ensure_ascii=False)}\n\n"

            # 执行审查（同步调用，在 SSE 生成器内）
            from ..db.database import async_session_factory

            async with async_session_factory() as review_db:
                service = ReviewService(review_db)
                results = await service.execute_review(
                    task_id=task_id,
                    dimensions=dimensions,
                    scope=request.scope,
                    chapter_ids=request.chapter_ids,
                    model_name=request.model_name,
                    provider_config_id=request.provider_config_id,
                    use_knowledge=request.use_knowledge,
                    knowledge_ids=request.knowledge_ids,
                )

            # 逐维度发送结果
            for dim in dimensions:
                label = dimension_labels.get(dim, dim)
                data = results.get(dim)
                if data:
                    yield f"data: {json.dumps({'type': 'result', 'dimension': dim, 'data': data}, ensure_ascii=False)}\n\n"
                yield f"data: {json.dumps({'type': 'progress', 'dimension': dim, 'status': 'completed', 'message': f'{label}审查完成'}, ensure_ascii=False)}\n\n"

            # 发送汇总
            summary = results.get("summary", {})
            yield f"data: {json.dumps({'type': 'summary', 'data': summary}, ensure_ascii=False)}\n\n"

            yield "data: [DONE]\n\n"

        except Exception as e:
            yield f"data: {json.dumps({'type': 'error', 'message': f'审查执行失败: {str(e)}'}, ensure_ascii=False)}\n\n"
            # 更新任务状态为失败
            try:
                from ..db.database import async_session_factory
                async with async_session_factory() as err_db:
                    err_task = await err_db.get(BidReviewTask, task_id)
                    if err_task:
                        err_task.status = ReviewTaskStatus.FAILED
                        err_task.error_message = str(e)
                        await err_db.commit()
            except Exception as err_update_exc:
                print(f"[Review] 更新任务失败状态异常: {err_update_exc}")
            yield "data: [DONE]\n\n"

    return sse_response(generate())


# ============ 获取审查结果 ============


@router.get("/result/{task_id}", response_model=ReviewResultResponse)
async def get_review_result(
    task_id: uuid.UUID,
    current_user: Annotated[User, Depends(require_editor)] = None,
    db: Annotated[AsyncSession, Depends(get_db)] = None,
):
    """获取审查结果（用于页面刷新后恢复）"""
    task = await db.get(BidReviewTask, task_id)
    if not task:
        raise HTTPException(status_code=404, detail="审查任务不存在")

    await _verify_project_member(task.project_id, current_user.id, db)

    return ReviewResultResponse(
        task_id=task.id,
        project_id=task.project_id,
        status=task.status,
        config={
            "dimensions": task.dimensions,
            "scope": task.scope,
            "model_name": task.model_name,
        },
        summary=task.summary,
        responsiveness=task.responsiveness_result,
        compliance=task.compliance_result,
        consistency=task.consistency_result,
        error_message=task.error_message,
        created_at=task.created_at,
    )


# ============ 审查历史 ============


@router.get("/history/{project_id}", response_model=ReviewHistoryResponse)
async def get_review_history(
    project_id: uuid.UUID,
    current_user: Annotated[User, Depends(require_editor)] = None,
    db: Annotated[AsyncSession, Depends(get_db)] = None,
):
    """获取项目的审查历史列表"""
    await _verify_project_member(project_id, current_user.id, db)

    result = await db.execute(
        select(BidReviewTask)
        .where(BidReviewTask.project_id == project_id)
        .order_by(BidReviewTask.created_at.desc())
    )
    tasks = result.scalars().all()

    from ..schemas.review import ReviewHistoryItem
    items = [
        ReviewHistoryItem(
            task_id=task.id,
            status=task.status,
            bid_filename=task.bid_filename,
            summary=task.summary,
            model_name=task.model_name,
            created_at=task.created_at,
        )
        for task in tasks
    ]

    return ReviewHistoryResponse(items=items)


# ============ 导出带批注的 Word ============


@router.post("/export-word")
async def export_reviewed_word(
    request: ReviewExportRequest,
    current_user: Annotated[User, Depends(require_editor)] = None,
    db: Annotated[AsyncSession, Depends(get_db)] = None,
):
    """导出带批注的 Word 文档"""
    task = await db.get(BidReviewTask, request.task_id)
    if not task:
        raise HTTPException(status_code=404, detail="审查任务不存在")

    await _verify_project_member(task.project_id, current_user.id, db)

    if task.status != ReviewTaskStatus.COMPLETED:
        raise HTTPException(status_code=400, detail="审查任务尚未完成")

    if not task.bid_file_path or not os.path.exists(task.bid_file_path):
        raise HTTPException(status_code=400, detail="投标文件不存在")

    try:
        # 根据请求的维度筛选导出内容
        dimensions = request.dimensions
        responsiveness = task.responsiveness_result if "responsiveness" in dimensions else None
        compliance = task.compliance_result if "compliance" in dimensions else None
        consistency = task.consistency_result if "consistency" in dimensions else None

        word_service = WordCommentService()
        output_path = word_service.export_reviewed_document(
            bid_file_path=task.bid_file_path,
            responsiveness=responsiveness,
            compliance=compliance,
            consistency=consistency,
        )

        # 返回文件
        with open(output_path, "rb") as f:
            buffer = f.read()

        import io
        buffer_io = io.BytesIO(buffer)

        filename = os.path.basename(output_path)
        encoded_filename = quote(filename)
        content_disposition = f"attachment; filename*=UTF-8''{encoded_filename}"

        return StreamingResponse(
            buffer_io,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": content_disposition},
        )

    except FileNotFoundError as e:
        raise HTTPException(status_code=400, detail=str(e)) from e
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"导出 Word 失败: {str(e)}",
        ) from e


# ============ 审查问题 AI 修复（流式） ============


class ApplyFixRequest(BaseModel):
    """应用审查问题修复请求"""
    chapter_id: str = Field(..., description="章节标识（用于定位）")
    current_content: str = Field(..., description="当前章节内容")
    issue_ids: list[str] = Field(default_factory=list, description="要修复的问题 ID 列表，空表示使用全部问题")


from pydantic import BaseModel as BaseModel_


@router.post("/apply-fix-stream/{task_id}")
async def apply_review_fix_stream(
    task_id: uuid.UUID,
    request: ApplyFixRequest,
    current_user: Annotated[User, Depends(require_editor)] = None,
    db: Annotated[AsyncSession, Depends(get_db)] = None,
):
    """对审查发现的问题，用 AI 生成修复建议，流式返回修改后内容"""
    from ..services.openai_service import OpenAIService

    task = await db.get(BidReviewTask, task_id)
    if not task:
        raise HTTPException(status_code=404, detail="审查任务不存在")

    await _verify_project_member(task.project_id, current_user.id, db)

    if task.status != ReviewTaskStatus.COMPLETED:
        raise HTTPException(status_code=400, detail="审查任务尚未完成")

    # 从审查结果里提取所有问题作为 suggestions
    all_issues: list[str] = []
    for result_field in [task.responsiveness_result, task.compliance_result, task.consistency_result]:
        if not result_field:
            continue
        # 结果通常是 {"issues": [...]} 或 {"items": [...]}
        for key in ("issues", "items", "results", "details"):
            items = result_field.get(key, [])
            if items and isinstance(items, list):
                for it in items:
                    if isinstance(it, dict):
                        issue_id = str(it.get("id", ""))
                        if request.issue_ids and issue_id not in request.issue_ids:
                            continue
                        desc = it.get("description") or it.get("issue") or it.get("content") or ""
                        suggestion = it.get("suggestion") or it.get("fix") or it.get("recommendation") or ""
                        text = f"{desc}" + (f"；建议：{suggestion}" if suggestion else "")
                        if text.strip():
                            all_issues.append(text.strip())
                break

    if not all_issues:
        raise HTTPException(status_code=400, detail="未找到可修复的审查问题")

    chapter_title = request.chapter_id

    openai_service = OpenAIService(db=db)
    await openai_service._ensure_initialized()

    if not openai_service.api_key:
        raise HTTPException(status_code=400, detail="请先配置 AI 服务 API 密钥")

    async def generate():
        try:
            yield f"data: {json.dumps({'status': 'started', 'message': '正在根据审查问题生成修改内容...'}, ensure_ascii=False)}\n\n"
            full_content = ""
            async for chunk in openai_service.rewrite_chapter_with_suggestions_stream(
                chapter_title=chapter_title,
                chapter_content=request.current_content,
                suggestions=all_issues,
            ):
                full_content += chunk
                yield f"data: {json.dumps({'status': 'streaming', 'content': chunk, 'full_content': full_content}, ensure_ascii=False)}\n\n"
            yield f"data: {json.dumps({'status': 'completed', 'content': full_content}, ensure_ascii=False)}\n\n"
        except Exception as e:
            yield f"data: {json.dumps({'status': 'error', 'message': str(e)}, ensure_ascii=False)}\n\n"
        yield "data: [DONE]\n\n"

    return sse_response(generate())
