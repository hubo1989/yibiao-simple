"""文档处理相关API路由"""

import uuid
from dataclasses import dataclass, field
from typing import Annotated, Any

from fastapi import APIRouter, UploadFile, File, HTTPException, Depends, Form, status
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, Field
from sqlalchemy import select, and_, exists
from sqlalchemy.ext.asyncio import AsyncSession

from ..models.schemas import (
    FileUploadResponse,
    AnalysisRequest,
    AnalysisType,
    WordExportRequest,
    ProjectFileUploadResponse,
    ProjectAnalysisRequest,
)
from ..models.user import User
from ..models.project import Project, project_members
from ..models.material import ChapterMaterialBinding
from ..db.database import get_db
from ..services.file_service import FileService
from ..services.material_service import render_content_with_material_bindings
from ..services.openai_service import OpenAIService
from ..services.prompt_service import PromptService
from ..utils.sse import sse_response
from ..auth.dependencies import require_editor

import json
import io
import re
import docx
from docx.shared import Pt, Cm, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement
from urllib.parse import quote

router = APIRouter(prefix="/api/document", tags=["文档处理"])


# ============ GB/T 9704 标准排版常量 ============
# 字号对照表（磅值）
FONT_SIZE_CHU = Pt(42)    # 初号
FONT_SIZE_XIAO_CHU = Pt(36)  # 小初
FONT_SIZE_YI = Pt(26)     # 一号
FONT_SIZE_XIAO_YI = Pt(24)   # 小一
FONT_SIZE_ER = Pt(22)     # 二号
FONT_SIZE_XIAO_ER = Pt(18)   # 小二
FONT_SIZE_SAN = Pt(16)    # 三号
FONT_SIZE_XIAO_SAN = Pt(15)  # 小三
FONT_SIZE_SI = Pt(14)     # 四号
FONT_SIZE_XIAO_SI = Pt(12)   # 小四号
FONT_SIZE_WU = Pt(10.5)   # 五号
FONT_SIZE_XIAO_WU = Pt(9)    # 小五号

LINE_SPACING_FIXED = Pt(28)  # 固定行间距 28 磅（GB/T 9704 标准）

FONT_HEITI = "黑体"       # 标题字体
FONT_FANGSONG = "仿宋"    # 正文字体
FONT_SONGTI = "宋体"      # 备用字体


# ============ 模板配置解析 ============

@dataclass
class TemplateConfig:
    """从 format_config JSON 解析出的渲染参数"""
    # 字体
    body_font: str = "仿宋"
    body_size_pt: float = 12.0
    h1_font: str = "黑体"
    h1_size_pt: float = 16.0
    h2_font: str = "黑体"
    h2_size_pt: float = 14.0
    h3_font: str = "黑体"
    h3_size_pt: float = 12.0
    table_font: str = "仿宋"
    table_size_pt: float = 10.5
    # 间距
    line_spacing_pt: float = 28.0
    first_indent_chars: int = 2
    h1_before: float = 24.0
    h1_after: float = 12.0
    h2_before: float = 12.0
    h2_after: float = 6.0
    h3_before: float = 6.0
    h3_after: float = 3.0
    # 页边距（mm）
    margin_top_mm: float = 37.0
    margin_bottom_mm: float = 35.0
    margin_left_mm: float = 28.0
    margin_right_mm: float = 26.0
    # 页码/页眉
    page_number_format: str = "第X页 共Y页"
    header_text: str = "{project_name}"
    header_position: str = "center"
    # 封面
    show_cover: bool = True
    cover_title_font: str = "黑体"
    cover_title_size_pt: float = 22.0
    cover_subtitle: str = "投标技术文件"
    show_bidder_info: bool = True
    cover_fields: list = field(default_factory=lambda: ["投标人", "编制日期"])
    # 目录
    show_toc: bool = True
    toc_title: str = "目  录"
    toc_levels: int = 3


def _parse_template_config(format_config: dict[str, Any] | None) -> TemplateConfig:
    """将 format_config JSON 解析为 TemplateConfig dataclass"""
    if not format_config:
        return TemplateConfig()

    cfg = TemplateConfig()
    font = format_config.get("font", {})
    spacing = format_config.get("spacing", {})
    margin = format_config.get("margin", {})
    page = format_config.get("page", {})
    cover = format_config.get("cover", {})
    toc = format_config.get("toc", {})

    # 字体
    if font.get("body_font"):
        cfg.body_font = font["body_font"]
    if font.get("body_size"):
        cfg.body_size_pt = float(font["body_size"])
    if font.get("h1_font"):
        cfg.h1_font = font["h1_font"]
    if font.get("h1_size"):
        cfg.h1_size_pt = float(font["h1_size"])
    if font.get("h2_font"):
        cfg.h2_font = font["h2_font"]
    if font.get("h2_size"):
        cfg.h2_size_pt = float(font["h2_size"])
    if font.get("h3_font"):
        cfg.h3_font = font["h3_font"]
    if font.get("h3_size"):
        cfg.h3_size_pt = float(font["h3_size"])
    if font.get("table_font"):
        cfg.table_font = font["table_font"]
    if font.get("table_size"):
        cfg.table_size_pt = float(font["table_size"])

    # 间距
    if spacing.get("line_spacing_pt"):
        cfg.line_spacing_pt = float(spacing["line_spacing_pt"])
    if spacing.get("first_indent_chars") is not None:
        cfg.first_indent_chars = int(spacing["first_indent_chars"])
    if spacing.get("h1_before") is not None:
        cfg.h1_before = float(spacing["h1_before"])
    if spacing.get("h1_after") is not None:
        cfg.h1_after = float(spacing["h1_after"])
    if spacing.get("h2_before") is not None:
        cfg.h2_before = float(spacing["h2_before"])
    if spacing.get("h2_after") is not None:
        cfg.h2_after = float(spacing["h2_after"])
    if spacing.get("h3_before") is not None:
        cfg.h3_before = float(spacing["h3_before"])
    if spacing.get("h3_after") is not None:
        cfg.h3_after = float(spacing["h3_after"])

    # 页边距
    if margin.get("top") is not None:
        cfg.margin_top_mm = float(margin["top"])
    if margin.get("bottom") is not None:
        cfg.margin_bottom_mm = float(margin["bottom"])
    if margin.get("left") is not None:
        cfg.margin_left_mm = float(margin["left"])
    if margin.get("right") is not None:
        cfg.margin_right_mm = float(margin["right"])

    # 页码/页眉
    if page.get("page_number_format"):
        cfg.page_number_format = page["page_number_format"]
    if page.get("header_text"):
        cfg.header_text = page["header_text"]
    if page.get("header_position"):
        cfg.header_position = page["header_position"]

    # 封面
    if cover.get("show_cover") is not None:
        cfg.show_cover = bool(cover["show_cover"])
    if cover.get("title_font"):
        cfg.cover_title_font = cover["title_font"]
    if cover.get("title_size"):
        cfg.cover_title_size_pt = float(cover["title_size"])
    if cover.get("subtitle"):
        cfg.cover_subtitle = cover["subtitle"]
    if cover.get("show_bidder_info") is not None:
        cfg.show_bidder_info = bool(cover["show_bidder_info"])
    if cover.get("cover_fields") is not None:
        cfg.cover_fields = list(cover["cover_fields"])

    # 目录
    if toc.get("show_toc") is not None:
        cfg.show_toc = bool(toc["show_toc"])
    if toc.get("toc_title"):
        cfg.toc_title = toc["toc_title"]
    if toc.get("toc_levels") is not None:
        cfg.toc_levels = int(toc["toc_levels"])

    return cfg


async def _load_template_config(
    template_id: str | None,
    db: AsyncSession,
) -> TemplateConfig:
    """
    从数据库加载模板配置。
    - template_id=None → 返回默认 GB/T 9704 配置（向后兼容）
    - template_id 有值但不存在 → 返回默认配置（graceful fallback）
    """
    if not template_id:
        return TemplateConfig()

    try:
        tid = uuid.UUID(template_id)
    except ValueError:
        return TemplateConfig()

    from ..models.export_template import ExportTemplate
    tmpl = await db.get(ExportTemplate, tid)
    if not tmpl:
        return TemplateConfig()

    return _parse_template_config(tmpl.format_config)


def _set_run_font(run, *, font_name=FONT_FANGSONG, font_size=None, bold=False):
    """设置 run 的字体（中西文统一）"""
    run.font.name = font_name
    if font_size:
        run.font.size = font_size
    run.bold = bold
    # 确保 rPr 存在
    if run._element.rPr is None:
        run._element._add_rPr()
    rpr = run._element.rPr
    # 设置 eastAsia 字体
    if rpr.rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rpr.insert(0, rFonts)
    rpr.rFonts.set(qn("w:eastAsia"), font_name)


def set_run_font_simsun(run: docx.text.run.Run) -> None:
    """兼容旧调用：设置宋体"""
    _set_run_font(run, font_name=FONT_SONGTI)


def set_paragraph_font_simsun(paragraph: docx.text.paragraph.Paragraph) -> None:
    """将段落内所有 runs 字体设置为宋体"""
    for run in paragraph.runs:
        set_run_font_simsun(run)


def _set_paragraph_format(paragraph, *, line_spacing_pt=28, first_line_indent_chars=2, space_before_pt=0, space_after_pt=0, font_size_pt=12, cfg: "TemplateConfig | None" = None):
    """设置段落格式：固定行间距、首行缩进、段前段后（GB/T 9704 标准）"""
    # 如果传了 cfg，优先使用 cfg 里的默认值
    if cfg is not None:
        if line_spacing_pt == 28:  # 仍是硬编码默认值，说明调用方没有覆盖
            line_spacing_pt = cfg.line_spacing_pt
        if first_line_indent_chars == 2:
            first_line_indent_chars = cfg.first_indent_chars
        if font_size_pt == 12:
            font_size_pt = cfg.body_size_pt

    pf = paragraph.paragraph_format
    # 固定值行间距
    from docx.enum.text import WD_LINE_SPACING
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(line_spacing_pt)
    if first_line_indent_chars and first_line_indent_chars > 0:
        pf.first_line_indent = Pt(font_size_pt * first_line_indent_chars)
    else:
        pf.first_line_indent = Pt(0)
    if space_before_pt:
        pf.space_before = Pt(space_before_pt)
    else:
        pf.space_before = Pt(0)
    pf.space_after = Pt(space_after_pt)


def _add_toc(doc, cfg: "TemplateConfig | None" = None):
    """添加自动目录页（GB/T 9704 样式：前导符点线 + 页码右对齐 + 层级缩进）"""
    _cfg = cfg or TemplateConfig()
    toc_title_text = _cfg.toc_title
    h_font = _cfg.h1_font

    # 目录标题：标题字体小二号，居中
    toc_heading = doc.add_paragraph()
    toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = toc_heading.add_run(toc_title_text)
    _set_run_font(run, font_name=h_font, font_size=FONT_SIZE_XIAO_ER, bold=True)
    toc_heading.paragraph_format.space_after = Pt(12)

    # 创建 TOC 样式（TOC 1 / TOC 2 / TOC 3）带前导符和右对齐制表位
    _ensure_toc_styles(doc, cfg)

    # 插入 TOC 字段 —— Word 打开时按 Ctrl+A → F9 更新
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar1)

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    # \\o "1-3"：收录 1-3 级标题 \\h：超链接 \\z：隐藏页码标签 \\u：使用大纲级别
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
    run._r.append(instrText)

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    run._r.append(fldChar2)

    # 占位提示
    run2 = paragraph.add_run("（打开 Word 后按 Ctrl+A → F9 更新目录）")
    _set_run_font(run2, font_name=_cfg.body_font, font_size=FONT_SIZE_WU)
    run2.font.color.rgb = docx.shared.RGBColor(0x99, 0x99, 0x99)

    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run2._r.append(fldChar3)

    doc.add_page_break()


def _ensure_toc_styles(doc, cfg: "TemplateConfig | None" = None):
    """确保 TOC 1/2/3 样式存在且带前导符点线、页码右对齐制表位"""
    _cfg = cfg or TemplateConfig()
    body_font = _cfg.body_font

    section = doc.sections[0]
    # 计算右边界位置（页面宽度 - 左右页边距）
    page_width = section.page_width or Cm(21.0)
    left_margin = section.left_margin or Cm(3.17)
    right_margin = section.right_margin or Cm(3.17)
    tab_pos = page_width - left_margin - right_margin

    toc_configs = [
        ("TOC 1", FONT_SIZE_SI, Cm(0), True),     # 一级：四号，无缩进，加粗
        ("TOC 2", FONT_SIZE_XIAO_SI, Cm(0.75), False),  # 二级：小四，缩进 0.75cm
        ("TOC 3", FONT_SIZE_XIAO_SI, Cm(1.5), False),   # 三级：小四，缩进 1.5cm
    ]

    for style_name, font_size, indent, bold in toc_configs:
        try:
            if style_name in doc.styles:
                style = doc.styles[style_name]
            else:
                style = doc.styles.add_style(style_name, docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)

            style.font.name = body_font
            style.font.size = font_size
            style.font.bold = bold
            if style._element.rPr is None:
                style._element._add_rPr()
            style._element.rPr.rFonts.set(qn("w:eastAsia"), body_font)

            pf = style.paragraph_format
            pf.left_indent = indent
            pf.space_before = Pt(2)
            pf.space_after = Pt(2)

            # 添加右对齐制表位 + 前导符点线
            pPr = style._element.get_or_add_pPr()
            # 移除已有 tabs
            for old_tabs in pPr.findall(qn('w:tabs')):
                pPr.remove(old_tabs)
            tabs = OxmlElement('w:tabs')
            tab = OxmlElement('w:tab')
            tab.set(qn('w:val'), 'right')
            tab.set(qn('w:leader'), 'dot')
            tab.set(qn('w:pos'), str(int(tab_pos.emu / 635)))  # EMU → twips
            tabs.append(tab)
            pPr.append(tabs)

        except Exception:
            pass  # 样式设置失败不影响文档生成


def _setup_header_footer(doc, project_name: str, cfg: "TemplateConfig | None" = None):
    """设置页眉（项目名称）和页脚—— 支持模板配置"""
    _cfg = cfg or TemplateConfig()
    body_font = _cfg.body_font

    section = doc.sections[0]
    # A4 页边距（从模板配置读取，单位 mm）
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(_cfg.margin_top_mm / 10)
    section.bottom_margin = Cm(_cfg.margin_bottom_mm / 10)
    section.left_margin = Cm(_cfg.margin_left_mm / 10)
    section.right_margin = Cm(_cfg.margin_right_mm / 10)

    # 页眉文本（支持 {project_name} 占位符）
    header_text = _cfg.header_text.replace("{project_name}", project_name)

    # 页眉对齐方式
    _header_align_map = {
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }
    header_align = _header_align_map.get(_cfg.header_position, WD_ALIGN_PARAGRAPH.CENTER)

    # 页眉：项目名称，小五号，底部单线
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.text = ""
    run = hp.add_run(header_text)
    _set_run_font(run, font_name=body_font, font_size=FONT_SIZE_XIAO_WU)
    run.font.color.rgb = docx.shared.RGBColor(0x66, 0x66, 0x66)
    hp.alignment = header_align
    # 页眉下边线
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    hp._p.get_or_add_pPr().append(pBdr)

    # 页脚：根据 page_number_format 决定格式
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.text = ""
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _add_field(para, field_code, default_text="1"):
        """插入 Word 字段"""
        run = para.add_run()
        _set_run_font(run, font_name=body_font, font_size=FONT_SIZE_XIAO_WU)
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar1)
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = field_code
        run._r.append(instrText)
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        run._r.append(fldChar2)
        run2 = para.add_run(default_text)
        _set_run_font(run2, font_name=body_font, font_size=FONT_SIZE_XIAO_WU)
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        run2._r.append(fldChar3)

    # 根据 page_number_format 生成页码
    fmt = _cfg.page_number_format
    if fmt == "X":
        # 只有页码数字
        _add_field(fp, ' PAGE ')
    elif fmt == "X/Y":
        _add_field(fp, ' PAGE ')
        r = fp.add_run("/")
        _set_run_font(r, font_name=body_font, font_size=FONT_SIZE_XIAO_WU)
        _add_field(fp, ' NUMPAGES ')
    else:
        # 默认：第 X 页 共 Y 页
        r = fp.add_run("第 ")
        _set_run_font(r, font_name=body_font, font_size=FONT_SIZE_XIAO_WU)
        _add_field(fp, ' PAGE ')
        r = fp.add_run(" 页  共 ")
        _set_run_font(r, font_name=body_font, font_size=FONT_SIZE_XIAO_WU)
        _add_field(fp, ' NUMPAGES ')
        r = fp.add_run(" 页")
        _set_run_font(r, font_name=body_font, font_size=FONT_SIZE_XIAO_WU)


def _add_word_table(doc, rows_text: list[str], add_markdown_runs_fn, cfg: "TemplateConfig | None" = None):
    """将 Markdown 表格行渲染为真正的 Word 表格"""
    _cfg = cfg or TemplateConfig()
    if not rows_text:
        return

    # 解析行为单元格
    parsed_rows = []
    for row_text in rows_text:
        cells = [c.strip() for c in row_text.split("|") if c.strip()]
        if cells:
            parsed_rows.append(cells)

    if not parsed_rows:
        return

    max_cols = max(len(r) for r in parsed_rows)
    table = doc.add_table(rows=len(parsed_rows), cols=max_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for r_idx, row_cells in enumerate(parsed_rows):
        row = table.rows[r_idx]
        for c_idx, cell_text in enumerate(row_cells):
            if c_idx < max_cols:
                cell = row.cells[c_idx]
                # 清空默认段落
                cell.paragraphs[0].text = ""
                add_markdown_runs_fn(cell.paragraphs[0], cell_text)
                # 单元格字号：模板表格字体/字号
                table_font_size = Pt(_cfg.table_size_pt)
                for run in cell.paragraphs[0].runs:
                    _set_run_font(run, font_name=_cfg.table_font, font_size=table_font_size)
                cell.paragraphs[0].paragraph_format.space_after = Pt(2)

        # 首行加灰色底色（表头）
        if r_idx == 0:
            for c_idx in range(max_cols):
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F2F2"/>')
                row.cells[c_idx]._tc.get_or_add_tcPr().append(shading)

    # 表格后空一行
    doc.add_paragraph()


@router.post("/upload", response_model=FileUploadResponse)
async def upload_file(
    file: UploadFile = File(...),
    current_user: Annotated[User, Depends(require_editor)] = None,
):
    """上传文档文件并提取文本内容"""
    try:
        # 检查文件类型（Content-Type + magic bytes 双重校验）
        allowed_types = [
            "application/pdf",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ]

        if file.content_type not in allowed_types:
            return FileUploadResponse(
                success=False, message="不支持的文件类型，请上传PDF或Word文档"
            )

        # Magic bytes 校验：防止伪造 Content-Type
        header = await file.read(8)
        await file.seek(0)
        is_pdf = header[:5] == b"%PDF-"
        is_docx = header[:4] == b"PK\x03\x04"  # docx 本质是 ZIP
        if not (is_pdf or is_docx):
            return FileUploadResponse(
                success=False, message="文件内容与类型不匹配，请上传有效的PDF或Word文档"
            )

        # 处理文件并提取文本
        file_content = await FileService.process_uploaded_file(file)

        return FileUploadResponse(
            success=True,
            message=f"文件 {file.filename} 上传成功",
            file_content=file_content,
        )

    except Exception as e:
        return FileUploadResponse(success=False, message=f"文件处理失败: {str(e)}")


@router.post("/analyze-stream")
async def analyze_document_stream(
    request: AnalysisRequest,
    current_user: Annotated[User, Depends(require_editor)] = None,
    db: AsyncSession = Depends(get_db),
):
    """流式分析文档内容"""
    try:
        # 创建OpenAI服务实例，从数据库加载配置
        openai_service = OpenAIService(db=db)
        if request.provider_config_id:
            configured = await openai_service.use_config_by_id(request.provider_config_id)
            if not configured:
                raise HTTPException(status_code=404, detail="所选 Provider 配置不存在")
        else:
            await openai_service._ensure_initialized()

        if not openai_service.api_key:
            raise HTTPException(status_code=400, detail="请先配置OpenAI API密钥")

        # 如果前端传了模型名称，覆盖默认模型
        if request.model_name:
            openai_service.set_model(request.model_name)

        async def generate():
            # 使用 PromptService 获取提示词（无项目上下文，使用全局/内置）
            prompt_service = PromptService(db)
            scene_key = {
                AnalysisType.OVERVIEW: "doc_analysis_overview",
                AnalysisType.REQUIREMENTS: "doc_analysis_requirements",
                AnalysisType.MATERIAL_REQUIREMENTS: "doc_analysis_requirements",
            }[request.analysis_type]
            prompt, _ = await prompt_service.get_prompt(scene_key)
            system_prompt, user_template = PromptService.split_prompt(prompt)
            user_prompt = prompt_service.render_prompt(
                user_template, {"file_content": request.file_content}
            )

            messages = [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ]

            # 流式返回分析结果
            async for chunk in openai_service.stream_chat_completion(
                messages, temperature=0.3
            ):
                yield f"data: {json.dumps({'chunk': chunk}, ensure_ascii=False)}\n\n"

            # 发送结束信号
            yield "data: [DONE]\n\n"

        return sse_response(generate())

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"文档分析失败: {e}") from e


@router.post("/export-word")
async def export_word(
    request: WordExportRequest,
    current_user: Annotated[User, Depends(require_editor)] = None,
    db: Annotated[AsyncSession, Depends(get_db)] = None,
):
    """根据目录数据导出Word文档"""
    try:
        doc = docx.Document()
        from docx.enum.text import WD_LINE_SPACING

        # ============ 加载模板配置（必须在样式初始化之前） ============
        cfg = await _load_template_config(request.template_id, db)
        body_font = cfg.body_font
        body_font_size = Pt(cfg.body_size_pt)
        h1_font = cfg.h1_font
        h1_size = Pt(cfg.h1_size_pt)
        h2_font = cfg.h2_font
        h2_size = Pt(cfg.h2_size_pt)
        h3_font = cfg.h3_font
        h3_size = Pt(cfg.h3_size_pt)
        line_spacing_pt = cfg.line_spacing_pt

        # ============ 文档样式初始化（从模板配置读取字体/行距） ============
        try:
            styles = doc.styles
            # Normal（正文）
            if "Normal" in styles:
                style = styles["Normal"]
                style.font.name = body_font
                style.font.size = body_font_size
                style.font.bold = False
                if style._element.rPr is None:
                    style._element._add_rPr()
                style._element.rPr.rFonts.set(qn("w:eastAsia"), body_font)
                style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                style.paragraph_format.line_spacing = Pt(line_spacing_pt)

            # Heading 1
            if "Heading 1" in styles:
                h1 = styles["Heading 1"]
                h1.font.name = h1_font
                h1.font.size = h1_size
                h1.font.bold = True
                h1.font.color.rgb = docx.shared.RGBColor(0, 0, 0)
                if h1._element.rPr is None:
                    h1._element._add_rPr()
                h1._element.rPr.rFonts.set(qn("w:eastAsia"), h1_font)
                h1.paragraph_format.space_before = Pt(cfg.h1_before)
                h1.paragraph_format.space_after = Pt(cfg.h1_after)
                h1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                h1.paragraph_format.line_spacing = Pt(max(line_spacing_pt, cfg.h1_size_pt + 6))

            # Heading 2
            if "Heading 2" in styles:
                h2 = styles["Heading 2"]
                h2.font.name = h2_font
                h2.font.size = h2_size
                h2.font.bold = True
                h2.font.color.rgb = docx.shared.RGBColor(0, 0, 0)
                if h2._element.rPr is None:
                    h2._element._add_rPr()
                h2._element.rPr.rFonts.set(qn("w:eastAsia"), h2_font)
                h2.paragraph_format.space_before = Pt(cfg.h2_before)
                h2.paragraph_format.space_after = Pt(cfg.h2_after)
                h2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                h2.paragraph_format.line_spacing = Pt(max(line_spacing_pt, cfg.h2_size_pt + 4))

            # Heading 3
            if "Heading 3" in styles:
                h3 = styles["Heading 3"]
                h3.font.name = h3_font
                h3.font.size = h3_size
                h3.font.bold = True
                h3.font.color.rgb = docx.shared.RGBColor(0, 0, 0)
                if h3._element.rPr is None:
                    h3._element._add_rPr()
                h3._element.rPr.rFonts.set(qn("w:eastAsia"), h3_font)
                h3.paragraph_format.space_before = Pt(cfg.h3_before)
                h3.paragraph_format.space_after = Pt(cfg.h3_after)
                h3.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                h3.paragraph_format.line_spacing = Pt(max(line_spacing_pt, cfg.h3_size_pt + 2))
        except Exception:
            pass

        title = request.project_name or "投标技术文件"

        # 页眉页脚 + 页边距
        _setup_header_footer(doc, title, cfg)

        # ============ 封面页 ============
        if cfg.show_cover:
            for _ in range(4):
                doc.add_paragraph()

            # 主标题：封面标题字体/字号，居中
            cover_title_size = Pt(cfg.cover_title_size_pt)
            title_p = doc.add_paragraph()
            title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_p.add_run(title)
            _set_run_font(title_run, font_name=cfg.cover_title_font, font_size=cover_title_size, bold=True)
            title_p.paragraph_format.space_after = Pt(6)

            # 副标题：正文字体三号，居中
            sub_p = doc.add_paragraph()
            sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sub_run = sub_p.add_run(cfg.cover_subtitle)
            _set_run_font(sub_run, font_name=body_font, font_size=FONT_SIZE_SAN)
            sub_p.paragraph_format.space_after = Pt(48)

            # 封面信息表
            if cfg.show_bidder_info:
                import datetime
                cover_field_values = {
                    "投标人": "（请填写投标单位名称）",
                    "编制日期": datetime.date.today().strftime("%Y 年 %m 月 %d 日"),
                }
                for field_label in cfg.cover_fields:
                    value = cover_field_values.get(field_label, "")
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r1 = p.add_run(f"{field_label}：")
                    _set_run_font(r1, font_name=h1_font, font_size=FONT_SIZE_SI)
                    r2 = p.add_run(value)
                    _set_run_font(r2, font_name=body_font, font_size=FONT_SIZE_SI)
                    p.paragraph_format.space_before = Pt(6)
                    p.paragraph_format.space_after = Pt(6)

            # 封面后分页
            doc.add_page_break()

        # 自动目录页
        if cfg.show_toc:
            _add_toc(doc, cfg)

        # 项目概述
        if request.project_overview:
            heading = doc.add_heading("项目概述", level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            overview_p = doc.add_paragraph(request.project_overview)
            for run in overview_p.runs:
                _set_run_font(run, font_name=body_font, font_size=body_font_size)
            _set_paragraph_format(overview_p, cfg=cfg)

        # 简单的 Markdown 段落解析：支持标题、列表、表格和基础加粗/斜体
        def add_markdown_runs(para: docx.text.paragraph.Paragraph, text: str) -> None:
            """在指定段落中追加 markdown 文本的 runs"""
            pattern = r"(\*\*.*?\*\*|\*.*?\*|`.*?`)"
            parts = re.split(pattern, text)
            for part in parts:
                if not part:
                    continue
                run = para.add_run()
                # 加粗
                if part.startswith("**") and part.endswith("**") and len(part) > 4:
                    run.text = part[2:-2]
                    run.bold = True
                # 斜体
                elif part.startswith("*") and part.endswith("*") and len(part) > 2:
                    run.text = part[1:-1]
                    run.italic = True
                # 行内代码：这里只去掉反引号
                elif part.startswith("`") and part.endswith("`") and len(part) > 2:
                    run.text = part[1:-1]
                else:
                    run.text = part
                # 正文字体
                _set_run_font(run, font_name=body_font, font_size=body_font_size)

        def add_markdown_paragraph(text: str) -> None:
            """将一段 Markdown 文本解析为一个普通段落，保留加粗/斜体效果"""
            para = doc.add_paragraph()
            add_markdown_runs(para, text)
            _set_paragraph_format(para, cfg=cfg)

        def parse_markdown_blocks(content: str):
            """
            识别 Markdown 内容中的块级元素，返回结构化的 block 列表：
            - ('list', items)        items: [(kind, num_str, text), ...]
            - ('table', rows)        rows: [text, ...]
            - ('heading', level, text)
            - ('paragraph', text)
            """
            blocks = []
            lines = content.split("\n")
            i = 0
            while i < len(lines):
                line = lines[i].rstrip("\r").strip()
                if not line:
                    i += 1
                    continue

                # 列表项（有序/无序）
                if (
                    line.startswith("- ")
                    or line.startswith("* ")
                    or re.match(r"^\d+\.\s", line)
                ):
                    # items: (kind, number, text)
                    items = []
                    while i < len(lines):
                        raw = lines[i].rstrip("\r")
                        stripped = raw.strip()
                        # 无序列表
                        if stripped.startswith("- ") or stripped.startswith("* "):
                            text = re.sub(r"^[-*]\s+", "", stripped).strip()
                            if text:
                                items.append(("unordered", None, text))
                            i += 1
                            continue
                        # 有序列表（1. xxx）
                        m_num = re.match(r"^(\d+)\.\s+(.*)$", stripped)
                        if m_num:
                            num_str, text = m_num.groups()
                            text = text.strip()
                            if text:
                                items.append(("ordered", num_str, text))
                            i += 1
                            continue
                        break

                    if items:
                        blocks.append(("list", items))
                    continue

                # 表格（简化为每行一个段落，单元格用 | 分隔）
                if "|" in line:
                    rows = []
                    while i < len(lines):
                        raw = lines[i].rstrip("\r")
                        stripped = raw.strip()
                        if "|" in stripped:
                            # 跳过仅由 - 和 | 组成的分隔行
                            if not re.match(r"^\|?[-\s\|]+\|?$", stripped):
                                cells = [c.strip() for c in stripped.split("|")]
                                row_text = " | ".join([c for c in cells if c])
                                if row_text:
                                    rows.append(row_text)
                            i += 1
                        else:
                            break
                    if rows:
                        blocks.append(("table", rows))
                    continue

                # Markdown 标题（# / ## / ###）
                if line.startswith("#"):
                    m = re.match(r"^(#+)\s*(.*)$", line)
                    if m:
                        level_marks, title_text = m.groups()
                        level = min(len(level_marks), 3)
                        blocks.append(("heading", level, title_text.strip()))
                    i += 1
                    continue

                # 普通段落：合并连续的普通行
                para_lines = []
                while i < len(lines):
                    raw = lines[i].rstrip("\r")
                    stripped = raw.strip()
                    if (
                        stripped
                        and not stripped.startswith("-")
                        and not stripped.startswith("*")
                        and "|" not in stripped
                        and not stripped.startswith("#")
                    ):
                        para_lines.append(stripped)
                        i += 1
                    else:
                        break
                if para_lines:
                    text = " ".join(para_lines)
                    blocks.append(("paragraph", text))
                else:
                    i += 1

            return blocks

        def render_markdown_blocks(blocks) -> None:
            """将结构化的 Markdown blocks 渲染到文档"""
            for block in blocks:
                kind = block[0]
                if kind == "list":
                    items = block[1]
                    for item_kind, num_str, text in items:
                        p = doc.add_paragraph()
                        if item_kind == "unordered":
                            run = p.add_run("• ")
                            _set_run_font(run, font_name=body_font, font_size=body_font_size)
                        else:
                            prefix = f"{num_str}."
                            run = p.add_run(prefix + " ")
                            _set_run_font(run, font_name=body_font, font_size=body_font_size)
                        add_markdown_runs(p, text)
                        # 列表项：左缩进、无首行缩进
                        _set_paragraph_format(p, first_line_indent_chars=0, cfg=cfg)
                        p.paragraph_format.left_indent = Cm(0.75)
                elif kind == "table":
                    rows = block[1]
                    _add_word_table(doc, rows, add_markdown_runs, cfg)
                elif kind == "heading":
                    _, level, text = block
                    heading = doc.add_heading(text, level=min(level, 3))
                    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    # 标题字体已由样式定义，无需额外设置
                elif kind == "paragraph":
                    _, text = block
                    add_markdown_paragraph(text)

        def add_markdown_content(content: str) -> None:
            """解析并渲染 Markdown 文本到文档"""
            blocks = parse_markdown_blocks(content)
            render_markdown_blocks(blocks)

        binding_map: dict[str, dict] = {}
        if request.project_id:
            try:
                project_uuid = uuid.UUID(request.project_id)
            except ValueError as exc:
                raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="无效的项目ID格式") from exc
            await _verify_project_member(project_uuid, current_user.id, db)
            result = await db.execute(
                select(ChapterMaterialBinding).where(ChapterMaterialBinding.project_id == project_uuid)
            )
            bindings = result.scalars().all()
            for binding in bindings:
                await db.refresh(binding, attribute_names=["material_asset"])
                material = binding.material_asset
                binding_map[str(binding.id)] = {
                    "caption": binding.caption or (material.name if material else "素材附件"),
                    "material_name": material.name if material else "素材附件",
                    "display_mode": binding.display_mode.value if hasattr(binding.display_mode, "value") else str(binding.display_mode),
                    "file_path": material.file_path if material else None,
                    "preview_path": material.preview_path if material else None,
                    "file_type": material.file_type if material else None,
                }

        # 递归构建文档内容（章节和内容）
        def add_outline_items(items, level: int = 1):
            for item in items:
                # 章节标题：使用 Heading 样式（黑体，字号由样式定义）
                if level <= 3:
                    heading = doc.add_heading(f"{item.id} {item.title}", level=level)
                    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    # 标题字体已由 Heading 样式统一设置为黑体
                else:
                    # 四级及以下：标题字体小四号加粗
                    para = doc.add_paragraph()
                    run = para.add_run(f"{item.id} {item.title}")
                    _set_run_font(run, font_name=h1_font, font_size=body_font_size, bold=True)
                    para.paragraph_format.space_before = Pt(cfg.h3_before)
                    para.paragraph_format.space_after = Pt(cfg.h3_after)

                # 叶子节点内容
                if not item.children:
                    content = item.content or ""
                    if content.strip():
                        render_content_with_material_bindings(
                            document=doc,
                            content=content,
                            binding_map=binding_map,
                            add_text=add_markdown_content,
                        )
                else:
                    add_outline_items(item.children, level + 1)

        add_outline_items(request.outline)

        # 输出到内存并返回
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        filename = f"{request.project_name or '标书文档'}.docx"
        # 使用 RFC 5987 格式对文件名进行 URL 编码，避免非 ASCII 字符导致的编码错误
        encoded_filename = quote(filename)
        content_disposition = f"attachment; filename*=UTF-8''{encoded_filename}"
        headers = {"Content-Disposition": content_disposition}

        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers=headers,
        )
    except Exception as e:
        import traceback

        logger_msg = f"导出Word失败: {str(e)}"
        print(logger_msg)
        traceback.print_exc()
        raise HTTPException(status_code=500, detail="导出Word失败，请稍后重试")


@router.post("/export-pdf")
async def export_pdf(
    request: WordExportRequest,
    current_user: Annotated[User, Depends(require_editor)] = None,
    db: Annotated[AsyncSession, Depends(get_db)] = None,
):
    """先生成 Word 文档，再通过 LibreOffice 转换为 PDF"""
    import tempfile
    import subprocess
    import os

    # 复用 export_word 的文档生成逻辑
    word_response = await export_word(request, current_user, db)

    # 读取生成的 Word 文档内容
    word_bytes = b""
    async for chunk in word_response.body_iterator:
        if isinstance(chunk, bytes):
            word_bytes += chunk
        else:
            word_bytes += chunk.encode()

    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = os.path.join(tmpdir, "output.docx")
        pdf_path = os.path.join(tmpdir, "output.pdf")

        with open(docx_path, "wb") as f:
            f.write(word_bytes)

        # 使用 LibreOffice headless 转换
        try:
            result = subprocess.run(
                ["soffice", "--headless", "--convert-to", "pdf", "--outdir", tmpdir, docx_path],
                capture_output=True,
                text=True,
                timeout=60,
            )
            if result.returncode != 0:
                raise RuntimeError(f"LibreOffice 转换失败: {result.stderr}")
        except FileNotFoundError:
            # 回退到 docx2pdf
            try:
                from docx2pdf import convert
                convert(docx_path, pdf_path)
            except Exception as e2:
                raise HTTPException(status_code=500, detail=f"PDF 转换工具不可用: {e2}")

        if not os.path.exists(pdf_path):
            raise HTTPException(status_code=500, detail="PDF 转换失败，未生成输出文件")

        with open(pdf_path, "rb") as f:
            pdf_bytes = f.read()

    buffer = io.BytesIO(pdf_bytes)
    filename = f"{request.project_name or '标书文档'}.pdf"
    encoded_filename = quote(filename)
    content_disposition = f"attachment; filename*=UTF-8''{encoded_filename}"

    return StreamingResponse(
        buffer,
        media_type="application/pdf",
        headers={"Content-Disposition": content_disposition},
    )


# ============ 项目上下文版本的接口 ============


async def _verify_project_member(
    project_id: uuid.UUID,
    user_id: uuid.UUID,
    db: AsyncSession,
) -> Project:
    """验证用户是项目成员并返回项目"""
    # 检查用户是否是项目成员
    member_exists = await db.execute(
        select(
            exists().where(
                and_(
                    project_members.c.project_id == project_id,
                    project_members.c.user_id == user_id,
                )
            )
        )
    )
    if not member_exists.scalar():
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="项目不存在或您没有访问权限",
        )

    # 获取项目
    result = await db.execute(select(Project).where(Project.id == project_id))
    project = result.scalar_one_or_none()
    if not project:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="项目不存在",
        )
    return project


@router.post("/upload-to-project", response_model=ProjectFileUploadResponse)
async def upload_file_to_project(
    project_id: Annotated[str, Form(...)],
    file: UploadFile = File(...),
    current_user: Annotated[User, Depends(require_editor)] = None,
    db: Annotated[AsyncSession, Depends(get_db)] = None,
):
    """上传文档文件到指定项目，提取文本内容并保存到项目记录"""
    try:
        # 验证 project_id 格式
        try:
            project_uuid = uuid.UUID(project_id)
        except ValueError:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="无效的项目ID格式",
            )

        # 验证用户是项目成员
        project = await _verify_project_member(project_uuid, current_user.id, db)

        # 检查文件类型（Content-Type + magic bytes 双重校验）
        allowed_types = [
            "application/pdf",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ]

        if file.content_type not in allowed_types:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="不支持的文件类型，请上传PDF或Word文档",
            )

        # Magic bytes 校验：防止伪造 Content-Type
        header = await file.read(8)
        await file.seek(0)
        is_pdf = header[:5] == b"%PDF-"
        is_docx = header[:4] == b"PK\x03\x04"  # docx 本质是 ZIP
        if not (is_pdf or is_docx):
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="文件内容与类型不匹配，请上传有效的PDF或Word文档",
            )

        # 处理文件并提取文本
        file_content = await FileService.process_uploaded_file(file)

        # 保存文件内容到项目记录
        project.file_content = file_content
        await db.flush()
        await db.refresh(project)

        return ProjectFileUploadResponse(
            success=True,
            message=f"文件 {file.filename} 已上传并保存到项目",
            project_id=str(project.id),
            file_content_length=len(file_content),
        )

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"文件处理失败: {str(e)}",
        )


@router.post("/analyze-project-stream")
async def analyze_project_document_stream(
    request: ProjectAnalysisRequest,
    current_user: Annotated[User, Depends(require_editor)] = None,
    db: Annotated[AsyncSession, Depends(get_db)] = None,
):
    """流式分析项目文档内容，并将分析结果保存到项目记录"""
    from ..db.database import async_session_factory

    try:
        # 验证 project_id 格式
        try:
            project_uuid = uuid.UUID(request.project_id)
        except ValueError:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="无效的项目ID格式",
            )

        # 验证用户是项目成员并获取项目
        project = await _verify_project_member(project_uuid, current_user.id, db)

        # 检查项目是否有文件内容
        if not project.file_content:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="项目尚未上传文档，请先上传招标文件",
            )

        # 创建OpenAI服务实例，从数据库加载配置
        openai_service = OpenAIService(db=db, project_id=project_uuid)
        if request.provider_config_id:
            configured = await openai_service.use_config_by_id(request.provider_config_id)
            if not configured:
                raise HTTPException(status_code=404, detail="所选 Provider 配置不存在")
        else:
            await openai_service._ensure_initialized()

        if not openai_service.api_key:
            raise HTTPException(status_code=400, detail="请先配置OpenAI API密钥")

        # 如果前端传了模型名称，覆盖默认模型
        if request.model_name:
            openai_service.set_model(request.model_name)

        # 用于收集分析结果的变量
        collected_result = []

        # 保存分析结果的变量
        analysis_type = request.analysis_type
        project_id = project_uuid

        async def generate():
            nonlocal collected_result

            # 使用 PromptService 获取提示词
            prompt_service = PromptService(db)
            scene_key = {
                AnalysisType.OVERVIEW: "doc_analysis_overview",
                AnalysisType.REQUIREMENTS: "doc_analysis_requirements",
                AnalysisType.MATERIAL_REQUIREMENTS: "doc_analysis_requirements",
            }[analysis_type]
            prompt, _ = await prompt_service.get_prompt(scene_key, project_id)
            system_prompt, user_template = PromptService.split_prompt(prompt)
            user_prompt = prompt_service.render_prompt(
                user_template, {"file_content": project.file_content}
            )

            messages = [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ]

            # 流式返回分析结果，同时收集完整结果
            async for chunk in openai_service.stream_chat_completion(
                messages, temperature=0.3
            ):
                collected_result.append(chunk)
                yield f"data: {json.dumps({'chunk': chunk}, ensure_ascii=False)}\n\n"

            # 在生成器内部创建新的数据库会话保存结果
            full_result = "".join(collected_result)
            async with async_session_factory() as save_db:
                try:
                    # 重新查询项目
                    save_project = await save_db.get(Project, project_id)
                    if save_project:
                        if analysis_type == AnalysisType.OVERVIEW:
                            save_project.project_overview = full_result
                        else:
                            save_project.tech_requirements = full_result
                        await save_db.commit()
                except Exception as e:
                    print(f"保存分析结果失败: {e}")

            # 发送结束信号
            yield "data: [DONE]\n\n"

        return sse_response(generate())

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"文档分析失败: {str(e)}",
        )


class ProjectAnalysisResult(BaseModel):
    """项目分析结果响应"""

    project_id: str
    file_content: str | None = None
    project_overview: str | None = None
    tech_requirements: str | None = None
    file_content_length: int = Field(..., description="文件内容字符数")


@router.get("/project-analysis/{project_id}", response_model=ProjectAnalysisResult)
async def get_project_analysis(
    project_id: uuid.UUID,
    current_user: Annotated[User, Depends(require_editor)] = None,
    db: Annotated[AsyncSession, Depends(get_db)] = None,
):
    """获取项目的文档分析结果"""
    # 验证用户是项目成员并获取项目
    project = await _verify_project_member(project_id, current_user.id, db)

    return ProjectAnalysisResult(
        project_id=str(project.id),
        file_content=project.file_content,
        project_overview=project.project_overview,
        tech_requirements=project.tech_requirements,
        file_content_length=len(project.file_content) if project.file_content else 0,
    )


class SaveAnalysisRequest(BaseModel):
    """保存分析结果请求"""
    project_overview: str | None = None
    tech_requirements: str | None = None


@router.post("/save-analysis/{project_id}")
async def save_project_analysis(
    project_id: uuid.UUID,
    request: SaveAnalysisRequest,
    current_user: Annotated[User, Depends(require_editor)] = None,
    db: Annotated[AsyncSession, Depends(get_db)] = None,
):
    """手动保存项目的文档分析结果（用于下一步前的备份保存）"""
    # 验证用户是项目成员并获取项目
    project = await _verify_project_member(project_id, current_user.id, db)

    # 保存分析结果
    if request.project_overview is not None:
        project.project_overview = request.project_overview
    if request.tech_requirements is not None:
        project.tech_requirements = request.tech_requirements

    await db.commit()

    return {"success": True, "message": "分析结果已保存"}
